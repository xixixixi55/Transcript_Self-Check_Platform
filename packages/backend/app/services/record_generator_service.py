"""
Layer 21: BE_Services — 笔录文档生成服务

负责：
1. 模板优先：使用 template.docx + template_filler_service 填充
2. 回退方案：构建文档结构 (document_builder_service) + officecli batch
3. 返回文件路径供下载
"""

import json
import os
import shutil
import subprocess
import tempfile
from datetime import datetime
from collections.abc import Mapping

from ..repository.template_approval_repository import TemplateApprovalRepository
from ..repository.template_registry_repository import TemplateRegistryRepository
from .document_builder_service import build_record_document
from .template_filler_service import fill_template
from .legacy_report_projection_service import project_ordered_legacy_report
from .template_profile_service import require_registered_template

# 模板文件路径
_TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(
    os.path.dirname(os.path.dirname(__file__))))), "word_templates", "template.docx")

# 查找 officecli 完整路径（Windows: .cmd, Unix: 无扩展名）
_OFFICECLI = shutil.which("officecli") or shutil.which("officecli.cmd")
if not _OFFICECLI:
    raise RuntimeError("officecli 未安装或不在 PATH 中。请运行: npm install -g officecli")

# subprocess 公共参数
# encoding='utf-8': 中文 Windows 默认 GBK，officecli 输出 UTF-8
_SUBPROCESS_KWARGS = dict(capture_output=True, encoding="utf-8")

def _run_officecli(*args: str) -> subprocess.CompletedProcess:
    """调用 officecli。uvicorn 子进程的 PATH 可能不完整（缺少 npm 全局目录
    和 System32），因此：
    1. 使用 shutil.which 查找 officecli 绝对路径（含 .CMD 扩展名）
    2. 在子进程 env 中显式注入 System32 路径，确保 Windows 能通过 cmd.exe
       执行 .CMD 批处理文件
    3. encoding='utf-8' 处理 officecli 的 UTF-8 输出
    """
    env = os.environ.copy()
    # 确保 cmd.exe 所在目录在 PATH 中（.CMD 文件依赖 cmd.exe 执行）
    system32 = r"C:\Windows\System32"
    if system32 not in env.get("PATH", ""):
        env["PATH"] = system32 + ";" + env.get("PATH", "")
    return subprocess.run(
        [_OFFICECLI, *args],
        env=env,
        **_SUBPROCESS_KWARGS,
    )


def generate_docx(report: dict, photo_paths: list[str] = None, output_dir: str = None,
                  archive_manifest: Mapping | None = None,
                  template_ref: Mapping | None = None,
                  template_registry: TemplateRegistryRepository | None = None,
                  template_approvals: TemplateApprovalRepository | None = None) -> str:
    """
    生成检查笔录 .docx 文件。
    优先使用模板填充，模板不存在时回退到 officecli batch 方案。
    返回生成的 .docx 文件路径。
    """
    if output_dir is None:
        output_dir = tempfile.mkdtemp()

    template_path = _TEMPLATE_PATH
    template_fingerprint = None
    if template_ref is not None:
        if template_registry is None or template_approvals is None:
            raise ValueError("template registry dependencies are required")
        registered = require_registered_template(
            template_registry, template_approvals, template_ref,
        )
        template_path = registered["internal_locator"]
        template_fingerprint = registered["fingerprint"]

    report = project_ordered_legacy_report(report)

    os.makedirs(output_dir, exist_ok=True)

    # 生成安全的文件名
    doc_number = report.get("document_number", "").replace("/", "-").replace("\\", "-")
    safe_doc_number = doc_number.replace("〔", "[").replace("〕", "]") if doc_number else ""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{safe_doc_number or '检查笔录'}_{timestamp}.docx"
    filepath = os.path.join(output_dir, filename)

    # 优先使用模板填充
    if os.path.isfile(template_path):
        template_options = (
            {} if template_fingerprint is None
            else {"expected_template_fingerprint": template_fingerprint}
        )
        try:
            if archive_manifest is None:
                fill_template(
                    report, template_path, filepath, photo_paths or [],
                    **template_options,
                )
            else:
                fill_template(
                    report, template_path, filepath, photo_paths or [], archive_manifest,
                    **template_options,
                )
            if os.path.isfile(filepath) and os.path.getsize(filepath) > 0:
                return filepath
        except Exception as e:
            if archive_manifest is not None or template_ref is not None:
                raise
            # 模板填充失败时回退到 batch 方案
            import traceback
            traceback.print_exc()

    # 回退：officecli batch 方案
    return _generate_via_batch(report, filepath, photo_paths or [])


def _insert_photos(filepath: str, photo_paths: list[str]):
    """向已生成的 docx 中插入用户上传的照片（附件2区域）"""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from lxml import etree

    doc = Document(filepath)
    body = doc.element.body
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 找到"附件2："段落
    target_idx = None
    for pi, para in enumerate(doc.paragraphs):
        if "附件2：" in para.text:
            target_idx = pi
            break

    if target_idx is None:
        print("[WARN] Photo target paragraph not found, skipping photo insert")
        return

    # 在附件2段落之后依次插入图片
    body_children = list(body)
    para_el_pos = None
    for i, child in enumerate(body_children):
        if child.tag == '{' + w_ns + '}p':
            # 这是第几个段落
            pass
    # 使用更简单的方法：在 doc.paragraphs 的指定位置后创建新段落
    # python-docx 不支持直接插入段落，用 lxml 操作
    para_el = doc.paragraphs[target_idx]._element
    body_els = list(body)
    try:
        insert_idx = body_els.index(para_el)
    except ValueError:
        print("[WARN] Cannot find paragraph element in body")
        return

    for i, photo_path in enumerate(photo_paths):
        if not os.path.isfile(photo_path):
            continue
        # 创建新段落元素并插入图片
        new_p = etree.SubElement(body, '{' + w_ns + '}p')
        new_pPr = etree.SubElement(new_p, '{' + w_ns + '}pPr')
        new_jc = etree.SubElement(new_pPr, '{' + w_ns + '}jc')
        new_jc.set('{' + w_ns + '}val', 'center')
        new_r = etree.SubElement(new_p, '{' + w_ns + '}r')
        new_rPr = etree.SubElement(new_r, '{' + w_ns + '}rPr')
        new_rPr_set = etree.SubElement(new_rPr, '{' + w_ns + '}sz')
        new_rPr_set.set('{' + w_ns + '}val', '32')  # 16pt

        new_run = doc.paragraphs[target_idx].add_run()  # Won't work for insertion
        # Instead: add photo to a temporary paragraph, then move it
        # Use a temp doc to create the image paragraph
        from io import BytesIO
        temp_doc = Document()
        temp_para = temp_doc.add_paragraph()
        temp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        temp_run = temp_para.add_run()
        temp_run.add_picture(photo_path, width=Inches(6.67), height=Inches(5.0))
        # Move temp_para element to the main doc
        temp_para_element = temp_para._element
        body.insert(insert_idx + 1 + i, temp_para_element)
        print(f"[OK] Photo inserted: {os.path.basename(photo_path)}")

    doc.save(filepath)


def _generate_via_batch(report: dict, filepath: str, photo_paths: list[str]) -> str:
    """使用 officecli batch 命令生成文档（原有方案）"""
    # 1. 构建文档命令
    commands = build_record_document(report, photo_paths)

    # 2. 创建空白 docx
    result = _run_officecli("create", filepath)
    if result.returncode != 0:
        raise RuntimeError(f"officecli create 失败: stdout={result.stdout} stderr={result.stderr}")

    # 3. 批量写入内容
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".json", delete=False, encoding="utf-8"
    ) as tmp:
        json.dump(commands, tmp, ensure_ascii=False)
        tmp_path = tmp.name

    try:
        result = _run_officecli("batch", filepath, "--input", tmp_path)
        if result.returncode != 0:
            raise RuntimeError(f"officecli batch 失败: {result.stderr}")
    finally:
        os.unlink(tmp_path)

    save_result = _run_officecli("save", filepath)
    if save_result.returncode != 0:
        raise RuntimeError(f"officecli save failed: {save_result.stderr}")

    if not os.path.isfile(filepath) or os.path.getsize(filepath) == 0:
        raise RuntimeError("officecli 生成的 Word 文档为空")

    return filepath
