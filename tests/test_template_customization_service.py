"""Tests for allow-listed front-end template customization."""

from __future__ import annotations

import os
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "packages", "backend"))

from app.repository.workbench_errors import WorkbenchPersistenceError  # noqa: E402
from app.services.docx_package_service import (  # noqa: E402
    compute_ooxml_package_fingerprint, read_validated_docx_entries,
)
from app.services.template_customization_service import (  # noqa: E402
    customize_template, read_template_customization,
)
from app.services.template_profile_service import validate_current_template_profile  # noqa: E402
from app.services.template_filler_service import fill_template  # noqa: E402
from test_legacy_report_projection_service import _report  # noqa: E402


def test_customization_changes_only_allowlisted_typography_and_keeps_profile(tmp_path: Path):
    source = Path(__file__).parents[1] / "word_templates" / "template.docx"
    before = source.read_bytes()
    destination = tmp_path / "SYNTHETIC-derived.docx"

    customize_template(source, destination, {
        "document_title": "SYNTHETIC 定制检查笔录",
        "body_font": "宋体",
        "body_font_size": 15,
    })

    assert source.read_bytes() == before
    source_entries = dict(read_validated_docx_entries(source))
    destination_entries = dict(read_validated_docx_entries(destination))
    assert destination_entries.keys() == source_entries.keys()
    assert {
        name for name in source_entries if destination_entries[name] != source_entries[name]
    } == {"word/document.xml"}

    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    source_root = etree.fromstring(source_entries["word/document.xml"])
    destination_root = etree.fromstring(destination_entries["word/document.xml"])
    source_paragraphs = source_root.xpath("/w:document/w:body/w:p", namespaces=namespace)
    destination_paragraphs = destination_root.xpath("/w:document/w:body/w:p", namespaces=namespace)
    attachment_index = next(
        index for index, paragraph in enumerate(source_paragraphs)
        if "".join(paragraph.xpath(".//w:t/text()", namespaces=namespace)).startswith("附件：")
    )
    for source_paragraph, destination_paragraph in zip(
        source_paragraphs[attachment_index:], destination_paragraphs[attachment_index:], strict=True,
    ):
        assert _canonical(source_paragraph) == _canonical(destination_paragraph)
    assert [
        _canonical(node) for node in source_root.xpath("/w:document/w:body/w:tbl", namespaces=namespace)
    ] == [
        _canonical(node) for node in destination_root.xpath("/w:document/w:body/w:tbl", namespaces=namespace)
    ]

    document = Document(destination)
    assert "SYNTHETIC 定制检查笔录" in document.paragraphs[0].text
    assert "{{title}}" not in document.paragraphs[0].text
    body_run = next(run for paragraph in document.paragraphs[2:] for run in paragraph.runs if run.text)
    assert body_run.font.name == "黑体"
    body_run = next(run for run in document.paragraphs[3].runs if run.text)
    assert body_run.font.size.pt == 15
    assert body_run._element.get_or_add_rPr().get_or_add_rFonts().get(qn("w:eastAsia")) == "宋体"
    assert read_template_customization(destination) == {
        "document_title": "SYNTHETIC 定制检查笔录",
        "body_font": "宋体",
        "body_font_size": 15,
    }
    fingerprint = compute_ooxml_package_fingerprint(destination)
    validate_current_template_profile(str(destination), document, fingerprint)

    report = _report()
    report["title"] = "SYNTHETIC 案件标题不应覆盖模板"
    output = tmp_path / "SYNTHETIC-filled.docx"
    fill_template(report, str(destination), str(output))
    filled = Document(output)
    assert "SYNTHETIC 定制检查笔录" in filled.paragraphs[0].text
    assert "SYNTHETIC 案件标题不应覆盖模板" not in filled.paragraphs[0].text


def _canonical(node) -> bytes:
    return etree.tostring(node, method="c14n", with_comments=True)


@pytest.mark.parametrize("customization", [
    {"document_title": "SYNTHETIC", "body_font": "Arial", "body_font_size": 16},
    {"document_title": "SYNTHETIC", "body_font": "宋体", "body_font_size": 20},
    {"document_title": "{{title}}", "body_font": "宋体", "body_font_size": 16},
])
def test_customization_rejects_values_outside_allowlist(tmp_path: Path, customization):
    source = Path(__file__).parents[1] / "word_templates" / "template.docx"
    with pytest.raises(WorkbenchPersistenceError) as error:
        customize_template(source, tmp_path / "SYNTHETIC-invalid.docx", customization)
    assert error.value.code == "TEMPLATE_CUSTOMIZATION_INVALID"
