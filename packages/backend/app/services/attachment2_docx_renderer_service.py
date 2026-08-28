"""渲染固定 current-template-v1 附件二图片槽位。"""
from __future__ import annotations
import copy
from typing import Any, Sequence
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Emu
from lxml import etree
from .attachment2_image_service import (
    ATTACHMENT2_CAPTION_LINE_TWIPS,
    ATTACHMENT2_DUAL_GROUP_IMAGE_ROW_HEIGHT_TWIPS,
    ATTACHMENT2_DUAL_GROUP_SLOT_HEIGHT_EMU,
    Attachment2PhotoAsset,
    calculate_fixed_geometry,
)
from .attachment_plan_models_service import Attachment2PagePlan, AttachmentPlan
from .attachment_plan_service import AttachmentPlanError
from .docx_attachment_xml_service import (
    W_NS,
    clone_page_break,
    qn,
    set_paragraph_text,
    text_of,
)
from .template_profile_service import CurrentTemplateProfile, TemplateProfileError

_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def build_attachment2_drawing(
    doc: Any, asset: Attachment2PhotoAsset, profile: CurrentTemplateProfile,
    used_drawing_ids: set[int], slot_height_emu: int | None = None,
) -> Any:
    geometry = calculate_fixed_geometry(
        asset.width_px, asset.height_px,
        slot_width_emu=profile.attachment2_slot_width_emu,
        slot_height_emu=(
            profile.attachment2_slot_height_emu
            if slot_height_emu is None else slot_height_emu
        ),
    )
    temporary = Document()
    run = temporary.add_paragraph().add_run()
    run.add_picture(
        asset.path, width=Emu(geometry.render_width_emu),
        height=Emu(geometry.render_height_emu),
    )
    drawing = run._r.find(".//%s" % qn(_WP_NS, "inline"))
    if drawing is None:
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2图片绘图结构无效。")
    image_part = doc.part.package.get_or_add_image_part(asset.path)
    relationship_id = doc.part.relate_to(image_part, RT.IMAGE)
    for blip in drawing.findall(".//%s" % qn(_A_NS, "blip")):
        blip.set(qn(_REL_NS, "embed"), relationship_id)
    assign_drawing_ids(drawing, used_drawing_ids)
    wrapper = etree.Element(qn(W_NS, "drawing"))
    wrapper.append(drawing)
    return wrapper


def existing_drawing_ids(root: Any) -> set[int]:
    values: set[int] = set()
    for element in root.iter():
        if element.tag in {qn(_WP_NS, "docPr"), qn(_PIC_NS, "cNvPr")}:
            try:
                values.add(int(element.get("id", "0")))
            except ValueError:
                continue
    return values


def assign_drawing_ids(root: Any, used: set[int]) -> None:
    next_id = max(used or {0}) + 1
    for element in root.iter():
        if element.tag not in {qn(_WP_NS, "docPr"), qn(_PIC_NS, "cNvPr")}:
            continue
        while next_id in used:
            next_id += 1
        element.set("id", str(next_id))
        used.add(next_id)
        next_id += 1


def find_attachment2_paragraph(body: Any, anchor: str, exact: bool = False) -> Any | None:
    for element in body.findall("./%s" % qn(W_NS, "p")):
        value = text_of(element)
        if (value == anchor if exact else anchor in value):
            return element
    return None


def twips(emu: int) -> int:
    return round(emu / 635)


def append_fixed_table_spacer(
    table: Any, height_twips: int, width_twips: int, column_count: int,
) -> None:
    row = etree.SubElement(table, qn(W_NS, "tr"))
    height = etree.SubElement(etree.SubElement(row, qn(W_NS, "trPr")), qn(W_NS, "trHeight"))
    height.attrib.update({qn(W_NS, "val"): str(height_twips), qn(W_NS, "hRule"): "exact"})
    cell = etree.SubElement(row, qn(W_NS, "tc"))
    cell_pr = etree.SubElement(cell, qn(W_NS, "tcPr"))
    etree.SubElement(cell_pr, qn(W_NS, "tcW"), {
        qn(W_NS, "w"): str(width_twips), qn(W_NS, "type"): "dxa",
    })
    if column_count > 1:
        etree.SubElement(cell_pr, qn(W_NS, "gridSpan")).set(
            qn(W_NS, "val"), str(column_count),
        )
    paragraph = etree.SubElement(cell, qn(W_NS, "p"))
    paragraph_pr = etree.SubElement(paragraph, qn(W_NS, "pPr"))
    etree.SubElement(paragraph_pr, qn(W_NS, "spacing"), {
        qn(W_NS, "before"): "0", qn(W_NS, "after"): "0",
    })


def render_attachment2(
    doc: Any,
    plan: AttachmentPlan,
    profile: CurrentTemplateProfile,
    assets: Sequence[Attachment2PhotoAsset],
) -> None:
    render_attachment2_pages(
        doc,
        plan.attachment2_pages,
        plan.attachment2_state.photo_count,
        profile,
        assets,
    )
def render_attachment2_pages(
    doc: Any,
    pages: Sequence[Attachment2PagePlan],
    photo_count: int,
    profile: CurrentTemplateProfile,
    assets: Sequence[Attachment2PhotoAsset],
) -> None:
    """使用显式页面块替换模板的附件二区域。"""
    body = doc.element.body
    label2 = find_attachment2_paragraph(body, profile.attachment2_label, exact=True)
    label3 = find_attachment2_paragraph(body, profile.attachment3_label, exact=True)
    if label2 is None or label3 is None:
        raise TemplateProfileError("当前模板附件二或附件三锚点丢失。")
    children = list(body)
    start = children.index(label2)
    end = children.index(label3)
    caption = next(
        (element for element in children[start + 1:end]
         if element.tag == qn(W_NS, "p") and "照片" in text_of(element)),
        None,
    )
    if caption is None:
        raise TemplateProfileError("当前模板附件二图片说明锚点丢失。")
    if not pages:
        for element in children[start:end]:
            body.remove(element)
        return
    _validate_assets(pages, photo_count, assets)
    page_break_anchor = copy.deepcopy(label2)
    preserved_caption = copy.deepcopy(caption)
    for element in children[start:end]:
        body.remove(element)
    nodes: list[Any] = []
    used_drawing_ids = existing_drawing_ids(body)
    for page_index, page in enumerate(pages):
        page_break = label2 if page_index == 0 else clone_page_break(page_break_anchor)
        _set_attachment2_page_spacing(page_break, 0 if page.layout == "four_grid" else profile.attachment2_page_break_after_twips)
        nodes.append(page_break)
        nodes.append(_build_page_table(
            doc, page, assets, profile, used_drawing_ids,
            preserved_caption, page.inspection_result_material_numbers,
        ))
    for offset, node in enumerate(nodes):
        body.insert(start + offset, node)
def _build_page_table(
    doc: Any,
    page: Attachment2PagePlan,
    assets: Sequence[Attachment2PhotoAsset],
    profile: CurrentTemplateProfile,
    used_drawing_ids: set[int],
    caption_template: Any,
    captions: Sequence[str],
) -> Any:
    _validate_material_groups(page)
    grid = _page_grid(page)
    column_count = len(grid[0])
    is_dual_group = page.layout == "four_grid"
    image_row_height = (
        ATTACHMENT2_DUAL_GROUP_IMAGE_ROW_HEIGHT_TWIPS
        if is_dual_group else profile.attachment2_slot_row_height_twips
    )
    slot_height_emu = (
        ATTACHMENT2_DUAL_GROUP_SLOT_HEIGHT_EMU
        if is_dual_group else profile.attachment2_slot_height_emu
    )
    expected_columns = (
        profile.attachment2_two_image_table_columns
        if page.layout == "two_centered"
        else profile.attachment2_four_image_table_columns
    )
    if (column_count != expected_columns
            or len(page.images) > profile.attachment2_max_images_per_page
            or len(page.images) % profile.attachment2_pair_size):
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2页面槽位约束无效。")
    slot_width_twips = twips(profile.attachment2_slot_width_emu)
    table_width_twips = slot_width_twips * profile.attachment2_slot_columns
    table = etree.Element(qn(W_NS, "tbl"))
    tbl_pr = etree.SubElement(table, qn(W_NS, "tblPr"))
    tbl_w = etree.SubElement(tbl_pr, qn(W_NS, "tblW"))
    tbl_w.set(qn(W_NS, "w"), str(table_width_twips))
    tbl_w.set(qn(W_NS, "type"), "dxa")
    etree.SubElement(tbl_pr, qn(W_NS, "jc")).set(qn(W_NS, "val"), "center")
    borders = etree.SubElement(tbl_pr, qn(W_NS, "tblBorders"))
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        etree.SubElement(borders, qn(W_NS, edge)).set(qn(W_NS, "val"), "nil")
    etree.SubElement(tbl_pr, qn(W_NS, "tblLayout")).set(qn(W_NS, "type"), "fixed")
    cell_margins = etree.SubElement(tbl_pr, qn(W_NS, "tblCellMar"))
    for edge in ("top", "left", "bottom", "right"):
        etree.SubElement(cell_margins, qn(W_NS, edge), {qn(W_NS, "w"): "0", qn(W_NS, "type"): "dxa"})
    table_grid = etree.SubElement(table, qn(W_NS, "tblGrid"))
    grid_widths = [slot_width_twips] * column_count
    for width in grid_widths:
        etree.SubElement(table_grid, qn(W_NS, "gridCol")).set(
            qn(W_NS, "w"), str(width),
        )
    if len(captions) != len(page.images) // profile.attachment2_pair_size:
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2图片对缺少对应检材文字。")
    for row_index, row_groups in enumerate(grid):
        row = etree.SubElement(table, qn(W_NS, "tr"))
        row_pr = etree.SubElement(row, qn(W_NS, "trPr"))
        height = etree.SubElement(row_pr, qn(W_NS, "trHeight"))
        height.set(qn(W_NS, "val"), str(image_row_height))
        height.set(qn(W_NS, "hRule"), "exact")
        for cell_index, image_group in enumerate(row_groups):
            cell = etree.SubElement(row, qn(W_NS, "tc"))
            cell_pr = etree.SubElement(cell, qn(W_NS, "tcPr"))
            cell_width = grid_widths[cell_index]
            etree.SubElement(cell_pr, qn(W_NS, "tcW"), {
                qn(W_NS, "w"): str(cell_width),
                qn(W_NS, "type"): "dxa",
            })
            margins = etree.SubElement(cell_pr, qn(W_NS, "tcMar"))
            for edge in ("top", "left", "bottom", "right"):
                etree.SubElement(margins, qn(W_NS, edge), {
                    qn(W_NS, "w"): "0",
                    qn(W_NS, "type"): "dxa",
                })
            etree.SubElement(cell_pr, qn(W_NS, "vAlign")).set(qn(W_NS, "val"), "center")
            paragraph = etree.SubElement(cell, qn(W_NS, "p"))
            paragraph_pr = etree.SubElement(paragraph, qn(W_NS, "pPr"))
            etree.SubElement(paragraph_pr, qn(W_NS, "spacing"), {
                qn(W_NS, "before"): "0", qn(W_NS, "after"): "0",
            })
            etree.SubElement(paragraph_pr, qn(W_NS, "jc")).set(
                qn(W_NS, "val"), "center",
            )
            for image in image_group:
                asset = assets[image.sequence_number - 1]
                run = etree.SubElement(paragraph, qn(W_NS, "r"))
                run.append(build_attachment2_drawing(
                    doc, asset, profile, used_drawing_ids, slot_height_emu,
                ))
        caption = captions[row_index]
        caption_gap = profile.attachment2_group_gap_twips if is_dual_group and row_index < len(grid) - 1 else 0
        caption_row = etree.SubElement(table, qn(W_NS, "tr"))
        caption_height = etree.SubElement(etree.SubElement(caption_row, qn(W_NS, "trPr")), qn(W_NS, "trHeight"))
        caption_height.attrib.update({qn(W_NS, "val"): str(ATTACHMENT2_CAPTION_LINE_TWIPS), qn(W_NS, "hRule"): "exact"})
        caption_cell = etree.SubElement(caption_row, qn(W_NS, "tc"))
        caption_pr = etree.SubElement(caption_cell, qn(W_NS, "tcPr"))
        etree.SubElement(caption_pr, qn(W_NS, "tcW"), {
            qn(W_NS, "w"): str(sum(grid_widths)), qn(W_NS, "type"): "dxa",
        })
        if column_count > 1:
            etree.SubElement(caption_pr, qn(W_NS, "gridSpan")).set(
                qn(W_NS, "val"), str(column_count),
            )
        etree.SubElement(caption_pr, qn(W_NS, "vAlign")).set(qn(W_NS, "val"), "center")
        caption_node = copy.deepcopy(caption_template)
        caption_spacing = caption_node.find("./%s/%s" % (qn(W_NS, "pPr"), qn(W_NS, "spacing")))
        if caption_spacing is not None:
            caption_spacing.set(qn(W_NS, "before"), "0")
            caption_spacing.set(qn(W_NS, "after"), "0")
        set_paragraph_text(caption_node, f"检材{caption}照片")
        caption_cell.append(caption_node)
        if caption_gap:
            append_fixed_table_spacer(table, caption_gap, sum(grid_widths), column_count)
    return table
def _validate_material_groups(page: Attachment2PagePlan) -> None:
    """确保渲染器使用规划器分组而不重新配对。"""
    if not 1 <= len(page.material_groups) <= 2:
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2检材组数量约束无效。")
    flattened = tuple(
        image for group in page.material_groups for image in group.images
    )
    if flattened != page.images:
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2检材组与图片计划不一致。")
    expected_numbers = tuple(group.material_number for group in page.material_groups)
    if page.inspection_result_material_numbers != expected_numbers:
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2检查结果检材编号计划无效。")
    for group in page.material_groups:
        if (len(group.images) != 2
                or any(image.evidence_number != group.material_number for image in group.images)):
            raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2检材图片组必须固定为两张。")
def _page_grid(page: Attachment2PagePlan) -> list[list[tuple[Any, ...]]]:
    """将显式槽位转换为固定表格网格，绝不使用 Word 自动流式布局。"""
    by_slot = {image.slot: image for image in page.images}
    if page.layout == "two_centered":
        expected = ("left", "right")
        if len(page.images) != 2 or set(by_slot) != set(expected):
            raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2横向页面计划无效。")
        return [[(by_slot["left"],), (by_slot["right"],)]]
    if page.layout == "four_grid":
        expected = ("top-left", "top-right", "bottom-left", "bottom-right")
        if len(page.images) != 4 or set(by_slot) != set(expected):
            raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2四图页面计划无效。")
        return [
            [(by_slot["top-left"],), (by_slot["top-right"],)],
            [(by_slot["bottom-left"],), (by_slot["bottom-right"],)],
        ]
    raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2页面布局类型无效。")
def _validate_assets(
    pages: Sequence[Attachment2PagePlan],
    photo_count: int,
    assets: Sequence[Attachment2PhotoAsset],
) -> None:
    if len(assets) != photo_count:
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件二图片计划与有效图片数量不一致。")
    expected = [
        image for page in pages
        for image in page.images
    ]
    if [image.sequence_number for image in expected] != list(range(1, len(assets) + 1)):
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件二图片顺序计划无效。")
    for page in pages:
        _page_grid(page)
def _set_attachment2_page_spacing(paragraph: Any, after_twips: int) -> None:
    p_pr = paragraph.find("./%s" % qn(W_NS, "pPr"))
    if p_pr is None:
        raise TemplateProfileError("当前模板附件二分页锚点缺少段落属性。")
    spacing = p_pr.find("./%s" % qn(W_NS, "spacing"))
    if spacing is None:
        raise TemplateProfileError("当前模板附件二分页锚点缺少行距属性。")
    spacing.set(qn(W_NS, "after"), str(after_twips))
__all__ = ["render_attachment2", "render_attachment2_pages"]
