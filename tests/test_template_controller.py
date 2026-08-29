"""T019 公共模板 API 与正式生成集成测试。"""

from __future__ import annotations

import json
import os
import shutil
import sys
import threading
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from unittest.mock import patch

import pytest
from fastapi.testclient import TestClient
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "packages", "backend"))

from app.repository.workbench.workbench_database import WorkbenchDatabase  # noqa: E402
from app.repository.workbench.workbench_errors import WorkbenchPersistenceError  # noqa: E402
from app.repository.case.shared_defaults_repository import SharedDefaultsRepository  # noqa: E402
from app.repository.template.template_approval_repository import TemplateApprovalRepository  # noqa: E402
from app.repository.template.template_registry_repository import TemplateRegistryRepository  # noqa: E402
from app.services.template.template_profile_service import (  # noqa: E402
    CURRENT_TEMPLATE_PACKAGE_FINGERPRINT,
    CURRENT_TEMPLATE_VERSION,
    CURRENT_TEMPLATE_VALIDATION_RULE,
    LEGACY_TEMPLATE_PACKAGE_FINGERPRINT,
    LEGACY_TEMPLATE_VERSION,
    PREVIOUS_TEMPLATE_PACKAGE_FINGERPRINT,
    PREVIOUS_TEMPLATE_VERSION,
    TemplateProfileError,
)
from app.services.runtime.workbench_factory_service import build_workbench_services  # noqa: E402
from app.services.template import template_registry_service  # noqa: E402
from app.services.document.docx_package_service import compute_ooxml_package_fingerprint  # noqa: E402
from synthetic_report_builders import build_ordered_report  # noqa: E402

CASE_ID = "case-SYNTHETIC-template-api"
REFERENCE = {
    "template_id": "electronic-inspection-record", "version": CURRENT_TEMPLATE_VERSION,
}
LEGACY_REFERENCE = {
    "template_id": "electronic-inspection-record", "version": LEGACY_TEMPLATE_VERSION,
}
PREVIOUS_REFERENCE = {
    "template_id": "electronic-inspection-record", "version": PREVIOUS_TEMPLATE_VERSION,
}
IDENTITY = {
    "identity_kind": "local_session",
    "client_instance_id": "SYNTHETIC-TEMPLATE-CLIENT",
    "session_id": "SYNTHETIC-TEMPLATE-SESSION",
    "deployment_instance_id": "SYNTHETIC-TEMPLATE-API",
}


@pytest.fixture()
def template_api(tmp_path: Path):
    from app.main import app
    from app.controllers import record_template_context_controller, template_controller

    database = WorkbenchDatabase(
        tmp_path / "workbench.sqlite3", IDENTITY["deployment_instance_id"],
    )
    services = build_workbench_services(database)
    report = build_ordered_report()
    report["inspection"].pop("primary_software", None)
    with database.transaction() as connection:
        connection.execute(
            "INSERT INTO case_shells(case_id,schema_version,case_number,case_name,case_summary,"
            "source_id,parse_task_id,lifecycle,report_available,revision,created_at,updated_at,deployment_instance_id) "
            "VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                CASE_ID, 1, "SYNTHETIC-001", "SYNTHETIC template API case",
                "SYNTHETIC summary", "source-SYNTHETIC-template-api",
                "task-SYNTHETIC-template-api", "review_ready", 1, 4,
                "2026-07-30T00:00:00+00:00", "2026-07-30T00:00:00+00:00", IDENTITY["deployment_instance_id"],
            ),
        )
        connection.execute(
            "INSERT INTO case_drafts VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                CASE_ID, 1, json.dumps(report), "legacy-v1", "{}", "[]", None,
                "plan-SYNTHETIC-stable", "review_ready", 4,
                "2026-07-30T00:00:00+00:00", "2026-07-30T00:00:00+00:00",
            ),
        )
    lease = services.leases.acquire(CASE_ID, IDENTITY)
    with patch.object(template_controller, "get_workbench_services", return_value=services), \
         patch.object(
             record_template_context_controller,
             "get_workbench_services",
             return_value=services,
         ):
        yield TestClient(app), services, lease


def _selection_body(lease: dict, revision: int = 4) -> dict:
    return {
        "template_ref": REFERENCE,
        "expected_revision": revision,
        "lease_id": lease["lease_id"],
        "lease_token": lease["lease_token"],
    }


def _derive_body(template_ref: dict, source_ref: dict | None = None) -> dict:
    return {
        "source_template_ref": source_ref or REFERENCE,
        "template_ref": template_ref,
        "display_name": "SYNTHETIC 前端微调模板",
        "customization": {
            "document_title": "SYNTHETIC 定制检查笔录",
            "body_font": "宋体",
            "body_font_size": 15,
        },
    }


def _export_report() -> dict:
    return {
        "title": "SYNTHETIC 电子数据检查笔录",
        "document_number": "SYN-TEST-019",
        "introduction": {"evidence_list": []},
        "inspection": {
            "primary_software": {
                "name": "SYNTHETIC 取证软件",
                "version": "V1.0",
                "confirmation_status": "confirmed_by_user",
            },
            "result": {},
        },
        "attachments": {"disc_number": "GP20260730-01"},
    }


def test_public_list_and_selection_are_safe_and_preserve_archive_facts(template_api):
    client, services, lease = template_api
    response = client.get("/api/v1/workbench/templates")
    assert response.status_code == 200, response.text
    templates = response.json()["data"]
    assert [item["template_ref"] for item in templates] == [REFERENCE]
    assert all(item["approval_record"]["status"] == "approved" for item in templates)
    assert "internal_locator" not in response.text
    assert str(services.database.database_path.parent) not in response.text

    with services.database.connect() as connection:
        before = {
            table: connection.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0]
            for table in ("task_records", "archive_attempts", "archive_plans", "archive_assets")
        }
    response = client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template",
        json=_selection_body(lease),
    )
    assert response.status_code == 200, response.text
    result = response.json()["data"]
    assert result["draft"]["template_ref"] == REFERENCE
    assert result["draft"]["revision"] == 5
    assert result["draft"]["archive_plan_id"] == "plan-SYNTHETIC-stable"
    assert result["impact"] == {
        "word_artifact_validity": "invalidated_by_template_change",
        "archive_plan_changed": False,
        "archive_task_created": False,
        "manifest_changed": False,
        "disc_mapping_changed": False,
    }
    with services.database.connect() as connection:
        after = {
            table: connection.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0]
            for table in before
        }
    assert after == before


def test_selection_rejects_client_governance_fields_lease_and_stale_revision(template_api):
    client, services, lease = template_api
    body = _selection_body(lease)
    body["fingerprint"] = "sha256:SYNTHETIC-forbidden"
    assert client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template", json=body,
    ).status_code == 422

    unknown = _selection_body(lease)
    unknown["template_ref"] = {
        "template_id": "template-SYNTHETIC-unknown", "version": "9.9.9",
    }
    response = client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template", json=unknown,
    )
    assert response.status_code == 404
    assert response.json()["detail"]["code"] == "TEMPLATE_UNKNOWN"

    historical = _selection_body(lease)
    historical["template_ref"] = LEGACY_REFERENCE
    response = client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template", json=historical,
    )
    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "HISTORICAL_TEMPLATE_READ_ONLY"

    historical["template_ref"] = PREVIOUS_REFERENCE
    response = client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template", json=historical,
    )
    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "HISTORICAL_TEMPLATE_READ_ONLY"

    assets = services.database.database_path.parent / "template-assets"
    assets.mkdir(exist_ok=True)
    pending_asset = assets / "SYNTHETIC-pending.docx"
    shutil.copy2(Path(__file__).parents[1] / "word_templates" / "template.docx", pending_asset)
    pending_ref = {"template_id": "template-SYNTHETIC-pending", "version": "1.0.0"}
    services.template_registry.register({
        "schema_version": 1,
        "template_ref": pending_ref,
        "display_name": "SYNTHETIC pending template",
        "fingerprint": CURRENT_TEMPLATE_PACKAGE_FINGERPRINT,
        "validation_rules": [CURRENT_TEMPLATE_VALIDATION_RULE],
        "asset_id": "asset-SYNTHETIC-pending",
        "registered_at": "2026-07-30T00:00:00+00:00",
    }, pending_asset)
    services.template_approvals.record(pending_ref, {
        "approval_record_id": "approval-SYNTHETIC-pending",
        "status": "pending",
        "acceptance_summary": "SYNTHETIC pending",
        "recorded_at": "2026-07-30T01:00:00+00:00",
    })
    pending = _selection_body(lease)
    pending["template_ref"] = pending_ref
    response = client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template", json=pending,
    )
    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "TEMPLATE_NOT_APPROVED"

    invalid_lease = _selection_body(lease)
    invalid_lease["lease_token"] = "token-SYNTHETIC-invalid"
    response = client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template", json=invalid_lease,
    )
    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "LEASE_NOT_ACTIVE"

    assert client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template",
        json=_selection_body(lease),
    ).status_code == 200
    response = client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template",
        json=_selection_body(lease),
    )
    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "REVISION_CONFLICT"
    assert str(client) not in response.text


def test_template_management_supports_upload_default_and_safe_revoke(template_api):
    client, services, _lease = template_api
    management_response = client.get("/api/v1/workbench/templates/management")
    assert management_response.status_code == 200, management_response.text
    management = management_response.json()["data"]
    assert management["default_template_ref"] == REFERENCE
    default_template = next(
        item for item in management["templates"] if item["template_ref"] == REFERENCE
    )
    assert default_template["is_default"] is True
    assert default_template["can_delete"] is False
    managed_references = {(
        item["template_ref"]["template_id"], item["template_ref"]["version"]
    ) for item in management["templates"]}
    assert (LEGACY_REFERENCE["template_id"], LEGACY_REFERENCE["version"]) not in managed_references
    assert (PREVIOUS_REFERENCE["template_id"], PREVIOUS_REFERENCE["version"]) not in managed_references
    with pytest.raises(WorkbenchPersistenceError):
        services.template_registry.get_internal(LEGACY_REFERENCE)
    with pytest.raises(WorkbenchPersistenceError):
        services.template_registry.get_internal(PREVIOUS_REFERENCE)
    historical_delete = client.delete(
        f"/api/v1/workbench/templates/electronic-inspection-record/{LEGACY_TEMPLATE_VERSION}",
    )
    assert historical_delete.status_code == 409
    assert historical_delete.json()["detail"]["code"] == "HISTORICAL_TEMPLATE_READ_ONLY"
    assert default_template["can_customize"] is True

    before_rename = services.template_registry.get_internal(REFERENCE)
    before_approval = services.template_approvals.require_approved(REFERENCE)
    before_default = services.defaults.get()["default_template_ref"]
    rename_response = client.put(
        f"/api/v1/workbench/templates/{REFERENCE['template_id']}/{REFERENCE['version']}/display-name",
        json={"display_name": "  SYNTHETIC 用户命名模版  "},
    )
    assert rename_response.status_code == 200, rename_response.text
    renamed = next(
        item for item in rename_response.json()["data"]["templates"]
        if item["template_ref"] == REFERENCE
    )
    assert renamed["display_name"] == "SYNTHETIC 用户命名模版"
    after_rename = services.template_registry.get_internal(REFERENCE)
    assert {
        key: after_rename[key] for key in after_rename if key != "display_name"
    } == {
        key: before_rename[key] for key in before_rename if key != "display_name"
    }
    assert services.template_approvals.require_approved(REFERENCE) == before_approval
    assert services.defaults.get()["default_template_ref"] == before_default

    for invalid_name in ("   ", "S" * 121):
        invalid = client.put(
            f"/api/v1/workbench/templates/{REFERENCE['template_id']}/{REFERENCE['version']}/display-name",
            json={"display_name": invalid_name},
        )
        assert invalid.status_code == 422
        assert invalid.json()["detail"]["code"] == "TEMPLATE_NAME_INVALID"
    extra = client.put(
        f"/api/v1/workbench/templates/{REFERENCE['template_id']}/{REFERENCE['version']}/display-name",
        json={"display_name": "SYNTHETIC", "unexpected": True},
    )
    assert extra.status_code == 422
    assert services.template_registry.get_internal(REFERENCE)["display_name"] == (
        "SYNTHETIC 用户命名模版"
    )

    source_template = Path(__file__).parents[1] / "word_templates" / "template.docx"
    upload_response = client.post(
        "/api/v1/workbench/templates",
        data={"display_name": "SYNTHETIC 上传笔录模版"},
        files={
            "file": (
                "SYNTHETIC-template.docx", source_template.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert upload_response.status_code == 200, upload_response.text
    uploaded_ref = upload_response.json()["data"]["template_ref"]
    assert uploaded_ref["template_id"].startswith("template-upload-")
    assert uploaded_ref["version"] == "1.0.0"
    assert upload_response.json()["data"]["template_ref"] == uploaded_ref

    second_upload = client.post(
        "/api/v1/workbench/templates",
        data={"display_name": "SYNTHETIC 第二个上传笔录模版"},
        files={
            "file": (
                "SYNTHETIC-template-2.docx", source_template.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert second_upload.status_code == 200, second_upload.text
    second_ref = second_upload.json()["data"]["template_ref"]
    assert second_ref["template_id"] != uploaded_ref["template_id"]
    assert second_ref["version"] == "1.0.0"

    updated = client.get("/api/v1/workbench/templates/management").json()["data"]
    assert any(item["template_ref"] == uploaded_ref for item in updated["templates"])
    with services.database.transaction() as connection:
        connection.execute(
            "UPDATE case_drafts SET template_ref_json = ? WHERE case_id = ?",
            (json.dumps(uploaded_ref), CASE_ID),
        )
    in_use_response = client.delete(
        f"/api/v1/workbench/templates/{uploaded_ref['template_id']}/{uploaded_ref['version']}",
    )
    assert in_use_response.status_code == 409
    assert in_use_response.json()["detail"]["code"] == "TEMPLATE_IN_USE"

    default_response = client.put(
        "/api/v1/workbench/templates/default",
        json={
            "template_ref": uploaded_ref,
            "expected_defaults_revision": updated["defaults_revision"],
        },
    )
    assert default_response.status_code == 200, default_response.text
    assert default_response.json()["data"]["default_template_ref"] == uploaded_ref

    revoke_response = client.delete(
        f"/api/v1/workbench/templates/electronic-inspection-record/{CURRENT_TEMPLATE_VERSION}",
    )
    assert revoke_response.status_code == 200, revoke_response.text
    remaining = revoke_response.json()["data"]
    assert all(item["template_ref"] != REFERENCE for item in remaining["templates"])

    blocked_response = client.delete(
        f"/api/v1/workbench/templates/{uploaded_ref['template_id']}/{uploaded_ref['version']}",
    )
    assert blocked_response.status_code == 409
    assert blocked_response.json()["detail"]["code"] == "DEFAULT_TEMPLATE_CANNOT_DELETE"
    assert services.defaults.get()["default_template_ref"] == uploaded_ref


def test_upload_rejects_invalid_template_layout(template_api, tmp_path: Path):
    client, _services, _lease = template_api
    source = Path(__file__).parents[1] / "word_templates" / "template.docx"
    broken = tmp_path / "SYNTHETIC-invalid-layout.docx"
    document = Document(str(source))
    document.paragraphs[0].text = ""
    document.save(broken)

    response = client.post(
        "/api/v1/workbench/templates",
        data={"display_name": "SYNTHETIC 旧版式上传模板"},
        files={
            "file": (
                "SYNTHETIC-invalid-layout.docx", broken.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "TEMPLATE_RULE_VALIDATION_FAILED"


def test_frontend_customization_derives_valid_immutable_template(template_api):
    client, services, _lease = template_api
    source = services.template_registry.get_internal(REFERENCE)
    source_bytes = Path(source["internal_locator"]).read_bytes()
    derived_ref = {
        "template_id": "template-SYNTHETIC-derived", "version": "1.1.0",
    }
    body = _derive_body(derived_ref)
    with services.database.transaction() as connection:
        connection.execute(
            "UPDATE case_drafts SET template_ref_json = ? WHERE case_id = ?",
            (json.dumps(REFERENCE), CASE_ID),
        )

    response = client.post("/api/v1/workbench/templates/derive", json=body)

    assert response.status_code == 200, response.text
    assert response.json()["data"]["template_ref"] == derived_ref
    assert services.templates.validate(derived_ref)["valid"] is True
    derived = services.template_registry.get_internal(derived_ref)
    assert Path(source["internal_locator"]).read_bytes() == source_bytes
    assert Path(derived["internal_locator"]).read_bytes() != source_bytes
    assert "SYNTHETIC 定制检查笔录" in Document(derived["internal_locator"]).paragraphs[0].text
    with services.database.connect() as connection:
        case_reference = json.loads(connection.execute(
            "SELECT template_ref_json FROM case_drafts WHERE case_id = ?", (CASE_ID,),
        ).fetchone()["template_ref_json"])
    assert case_reference == REFERENCE

    duplicate = client.post("/api/v1/workbench/templates/derive", json=body)
    assert duplicate.status_code == 409
    assert duplicate.json()["detail"]["code"] == "TEMPLATE_VERSION_IMMUTABLE"


def test_frontend_customization_rolls_back_registration_when_approval_fails(template_api):
    client, services, _lease = template_api
    derived_ref = {
        "template_id": "template-SYNTHETIC-approval-failure", "version": "1.0.0",
    }
    asset_root = services.database.database_path.parent / "template-assets"
    before = set(asset_root.glob("derived-template-*.docx"))

    with patch.object(
        services.template_approvals,
        "record",
        side_effect=WorkbenchPersistenceError("TEMPLATE_APPROVAL_CREATE_FAILED"),
    ):
        response = client.post(
            "/api/v1/workbench/templates/derive", json=_derive_body(derived_ref),
        )

    assert response.status_code == 500
    assert services.template_registry.find_internal(derived_ref) is None
    assert set(asset_root.glob("derived-template-*.docx")) == before


def test_frontend_customization_validation_failure_leaves_no_asset_or_record(template_api):
    client, services, _lease = template_api
    derived_ref = {
        "template_id": "template-SYNTHETIC-validation-failure", "version": "1.0.0",
    }
    asset_root = services.database.database_path.parent / "template-assets"
    before = set(asset_root.glob("derived-template-*.docx"))

    with patch.object(
        template_registry_service,
        "validate_current_template_profile",
        side_effect=WorkbenchPersistenceError("TEMPLATE_RULE_VALIDATION_FAILED"),
    ):
        response = client.post(
            "/api/v1/workbench/templates/derive", json=_derive_body(derived_ref),
        )

    assert response.status_code == 422
    assert services.template_registry.find_internal(derived_ref) is None
    assert set(asset_root.glob("derived-template-*.docx")) == before


def test_frontend_customization_rejects_unapproved_source(template_api):
    client, services, _lease = template_api
    asset_root = services.database.database_path.parent / "template-assets"
    asset_root.mkdir(exist_ok=True)
    pending_asset = asset_root / "SYNTHETIC-pending-derive.docx"
    shutil.copy2(Path(__file__).parents[1] / "word_templates" / "template.docx", pending_asset)
    pending_ref = {"template_id": "template-SYNTHETIC-pending-derive", "version": "1.0.0"}
    services.template_registry.register({
        "schema_version": 1,
        "template_ref": pending_ref,
        "display_name": "SYNTHETIC pending derive source",
        "fingerprint": CURRENT_TEMPLATE_PACKAGE_FINGERPRINT,
        "validation_rules": [CURRENT_TEMPLATE_VALIDATION_RULE],
        "asset_id": "asset-SYNTHETIC-pending-derive",
        "registered_at": "2026-07-30T00:00:00+00:00",
    }, pending_asset)
    services.template_approvals.record(pending_ref, {
        "approval_record_id": "approval-SYNTHETIC-pending-derive",
        "status": "pending",
        "acceptance_summary": "SYNTHETIC pending derive source",
        "recorded_at": "2026-07-30T01:00:00+00:00",
    })
    derived_ref = {
        "template_id": "template-SYNTHETIC-from-pending", "version": "1.0.0",
    }

    response = client.post(
        "/api/v1/workbench/templates/derive",
        json=_derive_body(derived_ref, pending_ref),
    )

    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "TEMPLATE_NOT_APPROVED"
    assert services.template_registry.find_internal(derived_ref) is None


def test_broken_title_slot_is_rejected_without_upload_or_derive_residue(template_api, tmp_path):
    client, services, _lease = template_api
    asset_root = services.database.database_path.parent / "template-assets"
    malformed = tmp_path / "SYNTHETIC-missing-title.docx"
    document = Document(Path(__file__).parents[1] / "word_templates" / "template.docx")
    for node in document.element.body[0].xpath(".//w:t"):
        node.text = ""
    document.save(malformed)
    before = set(asset_root.glob("*.docx"))

    with services.database.connect() as connection:
        registered_before = connection.execute(
            "SELECT COUNT(*) FROM template_versions",
        ).fetchone()[0]
    upload = client.post(
        "/api/v1/workbench/templates",
        data={"display_name": "SYNTHETIC broken title upload"},
        files={"file": (
            malformed.name, malformed.read_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )},
    )
    assert upload.status_code == 422
    with services.database.connect() as connection:
        registered_after = connection.execute(
            "SELECT COUNT(*) FROM template_versions",
        ).fetchone()[0]
    assert registered_after == registered_before
    assert set(asset_root.glob("*.docx")) == before

    source_asset = asset_root / "SYNTHETIC-broken-title-source.docx"
    shutil.copy2(malformed, source_asset)
    source_ref = {"template_id": "template-SYNTHETIC-broken-source", "version": "1.0.0"}
    services.template_registry.register({
        "schema_version": 1, "template_ref": source_ref,
        "display_name": "SYNTHETIC broken source",
        "fingerprint": compute_ooxml_package_fingerprint(source_asset),
        "validation_rules": [CURRENT_TEMPLATE_VALIDATION_RULE],
        "asset_id": "asset-SYNTHETIC-broken-source",
        "registered_at": "2026-07-30T00:00:00+00:00",
    }, source_asset)
    services.template_approvals.record(source_ref, {
        "approval_record_id": "approval-SYNTHETIC-broken-source",
        "status": "approved", "acceptance_summary": "SYNTHETIC legacy accepted source",
        "recorded_at": "2026-07-30T01:00:00+00:00",
    })
    derived_ref = {"template_id": "template-SYNTHETIC-broken-derived", "version": "1.0.0"}
    before_derive = set(asset_root.glob("*.docx"))
    derive = client.post(
        "/api/v1/workbench/templates/derive",
        json=_derive_body(derived_ref, source_ref),
    )
    assert derive.status_code == 422
    assert derive.json()["detail"]["code"] == "TEMPLATE_RULE_VALIDATION_FAILED"
    assert services.template_registry.find_internal(derived_ref) is None
    assert set(asset_root.glob("*.docx")) == before_derive


def test_concurrent_derive_same_reference_keeps_only_winning_asset(template_api):
    _client, services, _lease = template_api
    derived_ref = {
        "template_id": "template-SYNTHETIC-concurrent", "version": "1.0.0",
    }
    asset_root = services.database.database_path.parent / "template-assets"
    before = set(asset_root.glob("derived-template-*.docx"))
    barrier = threading.Barrier(2)
    original_customize = template_registry_service.customize_template

    def synchronized_customize(source, destination, customization):
        original_customize(source, destination, customization)
        barrier.wait(timeout=5)

    def derive_once():
        try:
            return services.templates.derive_customized(
                REFERENCE, derived_ref, "SYNTHETIC concurrent", _derive_body(derived_ref)["customization"],
            )
        except WorkbenchPersistenceError as error:
            return error.code

    with patch.object(
        template_registry_service, "customize_template", side_effect=synchronized_customize,
    ), ThreadPoolExecutor(max_workers=2) as executor:
        results = list(executor.map(lambda _index: derive_once(), range(2)))

    assert sum(isinstance(result, dict) for result in results) == 1
    assert results.count("TEMPLATE_VERSION_IMMUTABLE") == 1
    assert services.template_registry.find_internal(derived_ref) is not None
    assert len(set(asset_root.glob("derived-template-*.docx")) - before) == 1


def test_frontend_customization_rejects_readonly_invalid_and_extra_fields(template_api):
    client, services, _lease = template_api
    before = len(services.template_approvals.list_approved())
    base = {
        "source_template_ref": LEGACY_REFERENCE,
        "template_ref": {
            "template_id": "template-SYNTHETIC-invalid-derived", "version": "1.0.0",
        },
        "display_name": "SYNTHETIC invalid derived",
        "customization": {
            "document_title": "SYNTHETIC",
            "body_font": "宋体",
            "body_font_size": 16,
        },
    }
    readonly = client.post("/api/v1/workbench/templates/derive", json=base)
    assert readonly.status_code == 409
    assert readonly.json()["detail"]["code"] == "HISTORICAL_TEMPLATE_READ_ONLY"

    base["source_template_ref"] = REFERENCE
    base["customization"]["body_font"] = "Arial"
    invalid = client.post("/api/v1/workbench/templates/derive", json=base)
    assert invalid.status_code == 422
    assert invalid.json()["detail"]["code"] == "TEMPLATE_CUSTOMIZATION_INVALID"

    base["customization"]["body_font"] = "宋体"
    base["customization"]["unexpected"] = "SYNTHETIC"
    extra = client.post("/api/v1/workbench/templates/derive", json=base)
    assert extra.status_code == 422
    assert len(services.template_approvals.list_approved()) == before


def test_builtin_template_upgrade_migrates_legacy_cases_and_preserves_custom_default(tmp_path: Path):
    database = WorkbenchDatabase(tmp_path / "workbench.sqlite3", "SYNTHETIC-UPGRADE")
    template_root = Path(__file__).parents[1] / "word_templates"
    registry = TemplateRegistryRepository(database, (template_root,))
    approvals = TemplateApprovalRepository(database, registry)
    registry.register({
        "schema_version": 1, "template_ref": LEGACY_REFERENCE,
        "display_name": "电子数据检查笔录（current-template-v1）",
        "fingerprint": LEGACY_TEMPLATE_PACKAGE_FINGERPRINT,
        "validation_rules": [CURRENT_TEMPLATE_VALIDATION_RULE],
        "asset_id": "template-asset-current-v1",
        "registered_at": "2026-07-30T00:00:00+00:00",
    }, template_root / "template.docx")
    approvals.record(LEGACY_REFERENCE, {
        "approval_record_id": "template-approval-current-v1", "status": "approved",
        "acceptance_summary": "current-template-v1 已通过既有 Word、VML、分页、表格和附件验收。",
        "recorded_at": "2026-07-30T00:00:00+00:00",
    })
    SharedDefaultsRepository(database).ensure_default_template(LEGACY_REFERENCE)
    registry.register({
        "schema_version": 1, "template_ref": PREVIOUS_REFERENCE,
        "display_name": "电子数据检查笔录（current-template-v1）",
        "fingerprint": PREVIOUS_TEMPLATE_PACKAGE_FINGERPRINT,
        "validation_rules": [CURRENT_TEMPLATE_VALIDATION_RULE],
        "asset_id": "template-asset-current-v1-balanced",
        "registered_at": "2026-07-30T00:00:00+00:00",
    }, template_root / "template.docx")
    approvals.record(PREVIOUS_REFERENCE, {
        "approval_record_id": "template-approval-current-v1-balanced", "status": "approved",
        "acceptance_summary": "current-template-v1 已修正正文与附件一整体偏右并通过 Word 版式验收。",
        "recorded_at": "2026-07-30T00:00:00+00:00",
    })
    with database.transaction() as connection:
        connection.execute(
            "INSERT INTO case_shells(case_id,schema_version,case_number,case_name,case_summary,"
            "source_id,parse_task_id,lifecycle,report_available,revision,created_at,updated_at,"
            "deployment_instance_id) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                "case-SYNTHETIC-template-upgrade", 1, "SYNTHETIC-UPGRADE-001",
                "SYNTHETIC template upgrade", "SYNTHETIC summary",
                "source-SYNTHETIC-template-upgrade", None, "review_ready", 1, 1,
                "2026-08-13T00:00:00+00:00", "2026-08-13T00:00:00+00:00",
                database.deployment_instance_id,
            ),
        )
        connection.execute(
            "INSERT INTO case_drafts VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                "case-SYNTHETIC-template-upgrade", 1, json.dumps(build_ordered_report()),
                "legacy-v1", "{}", "[]", json.dumps(LEGACY_REFERENCE), None,
                "review_ready", 1, "2026-08-13T00:00:00+00:00",
                "2026-08-13T00:00:00+00:00",
            ),
        )
        connection.execute(
            "UPDATE template_versions SET internal_locator=? "
            "WHERE template_id=? AND version=?",
            (
                str(tmp_path / "SYNTHETIC-old-install" / "template.docx"),
                LEGACY_REFERENCE["template_id"], LEGACY_REFERENCE["version"],
            ),
        )
    upgraded = build_workbench_services(
        WorkbenchDatabase(database.database_path, database.deployment_instance_id),
    )
    assert upgraded.defaults.get()["default_template_ref"] == REFERENCE
    with pytest.raises(WorkbenchPersistenceError):
        upgraded.template_registry.get_internal(LEGACY_REFERENCE)
    with pytest.raises(WorkbenchPersistenceError):
        upgraded.template_registry.get_internal(PREVIOUS_REFERENCE)
    with upgraded.database.connect() as connection:
        saved_reference = json.loads(connection.execute(
            "SELECT template_ref_json FROM case_drafts WHERE case_id=?",
            ("case-SYNTHETIC-template-upgrade",),
        ).fetchone()["template_ref_json"])
    assert saved_reference == REFERENCE
    recovered = build_workbench_services(
        WorkbenchDatabase(database.database_path, database.deployment_instance_id),
    )
    assert recovered.defaults.get()["default_template_ref"] == REFERENCE

    custom_root = database.database_path.parent / "template-assets"
    custom_root.mkdir(exist_ok=True)
    custom_path = custom_root / "SYNTHETIC-custom.docx"
    shutil.copy2(Path(__file__).parents[1] / "word_templates" / "template.docx", custom_path)
    custom_ref = {"template_id": "template-SYNTHETIC-custom", "version": "1.0.0"}
    recovered.template_registry.register({
        "schema_version": 1, "template_ref": custom_ref,
        "display_name": "SYNTHETIC custom template",
        "fingerprint": CURRENT_TEMPLATE_PACKAGE_FINGERPRINT,
        "validation_rules": [CURRENT_TEMPLATE_VALIDATION_RULE],
        "asset_id": "asset-SYNTHETIC-custom",
        "registered_at": "2026-08-13T00:00:00+00:00",
    }, custom_path)
    recovered.template_approvals.record(custom_ref, {
        "approval_record_id": "approval-SYNTHETIC-custom", "status": "approved",
        "acceptance_summary": "SYNTHETIC custom approved",
        "recorded_at": "2026-08-13T00:00:00+00:00",
    })
    recovered.templates.set_default(custom_ref)

    restarted = build_workbench_services(
        WorkbenchDatabase(database.database_path, database.deployment_instance_id),
    )
    assert restarted.defaults.get()["default_template_ref"] == custom_ref


def test_builtin_template_rename_survives_service_restart(tmp_path: Path):
    database = WorkbenchDatabase(tmp_path / "workbench.sqlite3", "SYNTHETIC-RENAME-RESTART")
    services = build_workbench_services(database)
    renamed = "SYNTHETIC renamed built-in template"

    services.templates.rename_display_name(REFERENCE, renamed)
    restarted = build_workbench_services(
        WorkbenchDatabase(database.database_path, database.deployment_instance_id),
    )

    template = restarted.template_registry.get_internal(REFERENCE)
    assert template["display_name"] == renamed
    assert template["fingerprint"] == CURRENT_TEMPLATE_PACKAGE_FINGERPRINT
    assert restarted.template_approvals.require_approved(REFERENCE)["status"] == "approved"
    assert restarted.defaults.get()["default_template_ref"] == REFERENCE


def test_current_builtin_template_relocates_after_portable_directory_change(tmp_path: Path):
    database = WorkbenchDatabase(
        tmp_path / "workbench.sqlite3", "SYNTHETIC-CURRENT-TEMPLATE-RELOCATION",
    )
    services = build_workbench_services(database)
    with database.transaction() as connection:
        connection.execute(
            "UPDATE template_versions SET internal_locator=? "
            "WHERE template_id=? AND version=?",
            (
                str(tmp_path / "SYNTHETIC-old-portable" / "resources" /
                    "word_templates" / "template.docx"),
                REFERENCE["template_id"], REFERENCE["version"],
            ),
        )

    restarted = build_workbench_services(
        WorkbenchDatabase(database.database_path, database.deployment_instance_id),
    )

    current = restarted.template_registry.get_internal(REFERENCE)
    assert Path(current["internal_locator"]).resolve() == (
        Path(__file__).parents[1] / "word_templates" / "template.docx"
    ).resolve()
    assert current["fingerprint"] == CURRENT_TEMPLATE_PACKAGE_FINGERPRINT
    assert current["asset_id"] == "template-asset-current-v1-private-clean"
    assert restarted.templates.validate(REFERENCE)["valid"] is True


def test_builtin_template_upgrade_migrates_previous_default_and_case(tmp_path: Path):
    database = WorkbenchDatabase(tmp_path / "workbench.sqlite3", "SYNTHETIC-UPGRADE-PREVIOUS")
    template_root = Path(__file__).parents[1] / "word_templates"
    registry = TemplateRegistryRepository(database, (template_root,))
    approvals = TemplateApprovalRepository(database, registry)
    registry.register({
        "schema_version": 1, "template_ref": PREVIOUS_REFERENCE,
        "display_name": "电子数据检查笔录（current-template-v1）",
        "fingerprint": PREVIOUS_TEMPLATE_PACKAGE_FINGERPRINT,
        "validation_rules": [CURRENT_TEMPLATE_VALIDATION_RULE],
        "asset_id": "template-asset-current-v1-balanced",
        "registered_at": "2026-07-30T00:00:00+00:00",
    }, template_root / "template.docx")
    approvals.record(PREVIOUS_REFERENCE, {
        "approval_record_id": "template-approval-current-v1-balanced",
        "status": "approved",
        "acceptance_summary": "current-template-v1 已修正正文与附件一整体偏右并通过 Word 版式验收。",
        "recorded_at": "2026-07-30T00:00:00+00:00",
    })
    SharedDefaultsRepository(database).ensure_default_template(PREVIOUS_REFERENCE)
    case_id = "case-SYNTHETIC-previous-template-upgrade"
    with database.transaction() as connection:
        connection.execute(
            "INSERT INTO case_shells(case_id,schema_version,case_number,case_name,case_summary,"
            "source_id,parse_task_id,lifecycle,report_available,revision,created_at,updated_at,"
            "deployment_instance_id) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                case_id, 1, "SYNTHETIC-UPGRADE-PREVIOUS-001",
                "SYNTHETIC previous template upgrade", "SYNTHETIC summary",
                "source-SYNTHETIC-previous-template-upgrade", None,
                "review_ready", 1, 1, "2026-08-13T00:00:00+00:00",
                "2026-08-13T00:00:00+00:00", database.deployment_instance_id,
            ),
        )
        connection.execute(
            "INSERT INTO case_drafts VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                case_id, 1, json.dumps(build_ordered_report()), "legacy-v1", "{}", "[]",
                json.dumps(PREVIOUS_REFERENCE), None, "review_ready", 1,
                "2026-08-13T00:00:00+00:00", "2026-08-13T00:00:00+00:00",
            ),
        )
    upgraded = build_workbench_services(
        WorkbenchDatabase(database.database_path, database.deployment_instance_id),
    )
    assert upgraded.defaults.get()["default_template_ref"] == REFERENCE
    with pytest.raises(WorkbenchPersistenceError):
        upgraded.template_registry.get_internal(PREVIOUS_REFERENCE)
    with upgraded.database.connect() as connection:
        saved_reference = json.loads(connection.execute(
            "SELECT template_ref_json FROM case_drafts WHERE case_id=?", (case_id,),
        ).fetchone()["template_ref_json"])
    assert saved_reference == REFERENCE


def test_formal_export_resolves_server_reference_and_safely_revalidates(
    template_api, tmp_path: Path,
):
    from app.controllers import record_controller

    client, services, lease = template_api
    selection = client.put(
        f"/api/v1/workbench/cases/{CASE_ID}/template",
        json=_selection_body(lease),
    )
    revision = selection.json()["data"]["draft"]["revision"]
    output = tmp_path / "SYNTHETIC-T019.docx"
    output.write_bytes(b"SYNTHETIC-DOCX")
    with patch.object(record_controller, "generate_docx", return_value=str(output)) as generate:
        response = client.post(
            "/api/v1/records/export",
            data={
                "report_json": json.dumps(_export_report()),
                "case_id": CASE_ID,
                "case_revision": str(revision),
            },
        )
    assert response.status_code == 200
    assert generate.call_args.kwargs["template_ref"] == REFERENCE
    assert generate.call_args.kwargs["template_registry"] is services.template_registry
    assert generate.call_args.kwargs["template_approvals"] is services.template_approvals

    failure = TemplateProfileError(
        "所选模板指纹校验失败。", "TEMPLATE_FINGERPRINT_MISMATCH",
    )
    with patch.object(record_controller, "generate_docx", side_effect=failure):
        response = client.post(
            "/api/v1/records/export",
            data={
                "report_json": json.dumps(_export_report()),
                "case_id": CASE_ID,
                "case_revision": str(revision),
            },
        )
    assert response.status_code == 422
    assert response.json()["detail"] == {
        "code": "TEMPLATE_FINGERPRINT_MISMATCH",
        "message": "所选模板指纹校验失败。",
    }
    assert str(services.database.database_path.parent) not in response.text
