"""
Layer 21: BE_Services — 模板填充服务。
基于 template.docx，将 InspectionReport 数据填入 {{placeholder}} 占位符，
支持 {{#list}}...{{/list}} 列表块扩展。保留模板原始格式。

> 文件行数超过 250 行上限：本文件是模板填充的核心编排入口，包含 _flatten_report（报告
  扁平化）、_expand_all_lists（列表展开）、_expand_extract_table（表格展开）、
  _replace_placeholders（占位符替换）及多个辅助函数。拆分会增加参数传递复杂度。
"""
import copy
import os
import re
import tempfile
from collections.abc import Mapping
from typing import Any
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from .report_defaults_service import normalize_data_summary
from .entrust_person_service import format_entrust_persons
from .material_policy_service import (
    reviewed_material_display_name,
    material_from_legacy_item,
    select_display_identifiers,
)
from .attachment2_image_service import (
    EMU_PER_INCH,
    calculate_fixed_geometry,
    validate_attachment2_photos,
)
from .attachment2_plan_service import (
    build_attachment2_pages,
    material_photo_groups,
    photo_values,
)
from .attachment2_docx_renderer_service import render_attachment2_pages
from .attachment_plan_service import build_attachment_plan
from .attachment_docx_renderer_service import render_attachment_plan
from .docx_attachment_xml_service import (
    allow_latin_character_wrap,
    attachment1_source_lines,
)
from .docx_output_sanitizer_service import sanitize_generated_docx
from .template_profile_service import (
    CURRENT_TEMPLATE_PACKAGE_FINGERPRINT,
    current_template_profile,
    validate_current_template_profile,
    validate_template_package_fingerprint,
)
from .legacy_report_projection_service import project_ordered_legacy_report
from .hash_algorithm_service import hash_field_title, report_hash_algorithm

_ATTACHMENT_SUMMARY_GAP_TWIPS = 3 * 520


def fill_template(report: dict, template_path: str, output_path: str,
                  photo_paths: list[str] = None,
                  archive_manifest: Mapping | None = None,
                  expected_template_fingerprint: str | None = None,
                  template_ref: Mapping | None = None) -> str:
    """
    用 InspectionReport 数据填充模板，生成输出文档。
    返回 output_path。
    """
    if photo_paths is None:
        photo_paths = []
    report = project_ordered_legacy_report(report)

    plan = None
    profile = None
    photo_assets = ()
    expected_fingerprint = (
        expected_template_fingerprint or CURRENT_TEMPLATE_PACKAGE_FINGERPRINT
    )
    if archive_manifest is not None:
        raw_source_image_ids = (report.get("attachments") or {}).get("photo_ids") or []
        if not isinstance(raw_source_image_ids, list):
            raw_source_image_ids = []
        source_image_ids = tuple(
            str(value) for value in raw_source_image_ids
            if value is not None
        )
        if len(source_image_ids) != len(photo_paths):
            source_image_ids = tuple(
                f"photo-{index}" for index in range(1, len(photo_paths) + 1)
            )
        plan_report = copy.deepcopy(report)
        plan_report.setdefault("attachments", {})["photo_ids"] = list(source_image_ids)
        plan = build_attachment_plan(archive_manifest, plan_report)
        photo_assets = validate_attachment2_photos(photo_paths, source_image_ids)
        validate_template_package_fingerprint(template_path, expected_fingerprint)
    doc = Document(template_path)
    if archive_manifest is not None:
        profile = validate_current_template_profile(
            template_path, doc, expected_fingerprint, template_ref,
        )
    flat = _flatten_report(report)
    flat["photo_count"] = str(plan.attachment2_state.photo_count if plan else len(photo_paths))
    if plan is not None:
        flat["disc_number"] = plan.attachment_summary.disc_numbers[0]
        disc_date = _format_plan_date(plan.attachment_summary.inspection_date)
        flat["burning_date"] = disc_date
        flat["created_date"] = disc_date

    # 1. 展开列表块（必须先做，因为会复制段落）
    _expand_all_lists(doc, report)

    # 2. 展开表格提取清单
    if plan is None:
        _expand_extract_table(doc, report)

    # 3. 替换简单 {{key}} 占位符
    _replace_placeholders(doc, flat)
    _update_inspection_result(doc, report, flat, plan)

    # 3.5 替换页眉/页脚占位符
    _replace_header_footer(doc, flat)
    _enable_dynamic_page_fields(doc)

    # 3.6 固定模板附件区域必须先由可信 manifest 计划渲染，再替换剩余 VML 占位符。
    if plan is not None:
        render_attachment_plan(doc, plan, profile, report, photo_assets)
        _update_attachment_summary(doc, plan)
    _replace_vml_textbox_placeholders(doc, flat)

    # 4. 处理照片附件
    if plan is None:
        _handle_photos(doc, photo_paths, report)

    # 4.3 清理附件间多余空段落
    _cleanup_attachment_spacing(doc)

    # 4.4 法定标题格式在所有动态渲染完成后统一施加。
    _apply_required_heading_styles(doc)
    _apply_required_attachment_table_styles(doc)

    # 4.5 清除批注引用
    _remove_comments(doc)

    # 5. 保存
    output_path = os.fspath(output_path)
    output_dir = os.path.dirname(output_path) or None
    staged_fd, staged_path = tempfile.mkstemp(
        prefix=".docx-stage-", suffix=".docx", dir=output_dir
    )
    os.close(staged_fd)
    os.unlink(staged_path)
    try:
        doc.save(staged_path)
        sanitize_generated_docx(staged_path)
        os.replace(staged_path, output_path)
    finally:
        if os.path.exists(staged_path):
            os.unlink(staged_path)
    return output_path


def _format_plan_date(value: str) -> str:
    year, month, day = value.split("-")
    return f"{year}年{int(month)}月{int(day)}日"


def _update_attachment_summary(doc: Document, plan) -> None:
    """Write the storage-medium summary from the validated manifest plan."""
    attachment1_summary = (
        f"附件：1、电子数据提取固定清单，共{len(plan.attachment1_pages)}页；"
    )
    for paragraph in doc.paragraphs:
        if "1、电子数据提取固定清单" in paragraph.text:
            _replace_paragraph_text(paragraph, attachment1_summary)
            break
    photo_summary = (
        f"2、检材图{plan.attachment2_state.photo_count}张，"
        f"共{len(plan.attachment2_pages)}页；"
    )
    for paragraph in doc.paragraphs:
        if "2、检材图" in paragraph.text:
            _replace_paragraph_text(paragraph, photo_summary)
            break
    disc_numbers = plan.attachment_summary.disc_numbers
    count = len(plan.attachment3_pages)
    if plan.archive_medium == "hard_drive":
        summary = f"3、本鉴定中心拷贝的编号为“{disc_numbers[0]}”的硬盘1块，共1页。"
    else:
        joined = "、".join(disc_numbers)
        summary = (
            f"3、本鉴定中心刻制的编号为“{joined}”的光盘{count}张，"
            f"共{count}页。"
        )
    for paragraph in doc.paragraphs:
        if "3、本鉴定中心刻制的" in paragraph.text:
            nodes = paragraph._element.findall(
                ".//{%s}t" % "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            )
            if nodes:
                nodes[0].text = summary
                for node in nodes[1:]:
                    node.text = ""
            return


def _update_inspection_result(
    doc: Document, report: dict, flat: dict, plan=None,
) -> None:
    """Render the reviewed materials and trusted manifest part details."""
    evidence_numbers = _ordered_unique(
        item.get("evidence_number")
        for item in (report.get("introduction") or {}).get("evidence_list") or []
        if isinstance(item, Mapping)
    )
    if not evidence_numbers:
        return
    evidence_label = "、".join(evidence_numbers)
    primary = (report.get("inspection") or {}).get("primary_software")
    software_name = (
        str(primary.get("name")).strip()
        if isinstance(primary, Mapping) and primary.get("name")
        else flat["software_name"]
    )
    software_version = (
        str(primary.get("version")).strip()
        if isinstance(primary, Mapping) and primary.get("version")
        else flat["software_version"]
    )
    result_text = (
        f"经对编号为{evidence_label}号检材使用{software_name}（版本号为"
        f"{software_version}）进行检查，检出{flat['data_summary']}等电子数据。"
    )
    if plan is None:
        hash_title = hash_field_title(report_hash_algorithm(report))
        result_text += (
            f"将检出结果生成为“{flat['rar_filename']}”文件，"
            f"{hash_title}为“{flat['md5_hash'].upper()}”，"
            f"文件大小为“{flat['file_size']}”字节。"
        )
        if flat.get("disc_number"):
            result_text += f"结果以封盘方式刻录在编号为“{flat['disc_number']}”的光盘中。"
    else:
        parts = plan.attachment3_pages
        hash_title = hash_field_title(plan.hash_algorithm)
        part_text = "；".join(
            f"“{part.filename}”文件，{hash_title}为“{part.md5.upper()}”，"
            f"文件大小为“{part.size_bytes}”字节"
            for part in parts
        )
        disc_numbers = _ordered_unique(part.disc_number for part in parts)
        result_text += f"将检出结果生成为{part_text}。"
        if plan.archive_medium == "hard_drive":
            result_text += (
                f"结果以拷贝的方式保存在编号为“{disc_numbers[0]}”的硬盘中。"
            )
        else:
            result_text += f"结果以封盘方式刻录在编号为“{'、'.join(disc_numbers)}”的光盘中。"
    for paragraph in doc.paragraphs:
        if "经对编号为" in paragraph.text:
            _replace_paragraph_text(paragraph, result_text)
            return


def _ordered_unique(values) -> list[str]:
    result = []
    for value in values:
        text = "" if value is None else str(value).strip()
        if text and text not in result:
            result.append(text)
    return result


def _replace_paragraph_text(paragraph: Any, value: str) -> None:
    nodes = paragraph._element.findall(
        ".//{%s}t" % "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    )
    if not nodes:
        return
    nodes[0].text = value
    for node in nodes[1:]:
        node.text = ""


def _enable_dynamic_page_fields(doc: Document) -> None:
    """Ask Word to refresh PAGE/NUMPAGES fields when opening or printing."""
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    settings = doc.settings.element
    update_fields = settings.find("./{%s}updateFields" % w_ns)
    if update_fields is None:
        update_fields = settings.makeelement("{%s}updateFields" % w_ns)
        anchor = settings.find("./{%s}hdrShapeDefaults" % w_ns)
        if anchor is None:
            anchor = settings.find("./{%s}footnotePr" % w_ns)
        if anchor is None:
            anchor = settings.find("./{%s}compat" % w_ns)
        if anchor is None:
            settings.append(update_fields)
        else:
            settings.insert(settings.index(anchor), update_fields)
    update_fields.set("{%s}val" % w_ns, "true")
    for relationship in doc.part.rels.values():
        if relationship.reltype != RT.FOOTER:
            continue
        for field in relationship.target_part.element.findall(
                ".//{%s}fldChar" % w_ns):
            if field.get("{%s}fldCharType" % w_ns) == "begin":
                field.set("{%s}dirty" % w_ns, "true")


# ═══════════════════════════════════════════
# 报告字段扁平化
# ═══════════════════════════════════════════

def _flatten_report(report: dict) -> dict:
    """将嵌套 InspectionReport 扁平化为模板占位符键值对"""
    intro = report.get("introduction", {})
    insp = report.get("inspection", {})
    result = insp.get("result", {})
    attach = report.get("attachments", {})
    burning_date = attach.get("burning_date", "")
    evidence_list = intro.get("evidence_list", [])
    evidence_numbers = tuple(
        _ordered_unique(
            item.get("evidence_number")
            for item in evidence_list
            if isinstance(item, Mapping)
        )
    )

    # 软件工具合并格式文本
    tools = insp.get("software_tools", [])
    tool_parts = []
    for t in tools:
        name = t.get("name", "")
        ver = t.get("version", "")
        if name and ver:
            tool_parts.append(f"{name}（版本号为{ver}）")
        elif name:
            tool_parts.append(name)
    software_tools_text = "，".join(tool_parts)

    flat = {
        "title": report.get("title", ""),
        "document_number": report.get("document_number", ""),
        "entrust_unit": (
            str(intro.get("entrust_unit_prefix", "")).strip()
            + str(intro.get("entrust_unit", "")).strip()
        ),
        "entrust_persons_text": format_entrust_persons(intro.get("entrust_persons")),
        "entrust_time": intro.get("entrust_time", ""),
        "case_summary": intro.get("case_summary", ""),
        "inspection_requirement": intro.get("inspection_requirement", ""),
        "inspection_time_range": intro.get("inspection_time_range", ""),
        "inspection_place": intro.get("inspection_place", ""),
        "method": insp.get("method", ""),
        "hardware_device": insp.get("hardware_device", ""),
        "software_tools_text": software_tools_text,
        "evidence_number": "、".join(evidence_numbers) or result.get("evidence_number", ""),
        "software_name": result.get("software_name", ""),
        "software_version": result.get("software_version", ""),
        "data_summary": normalize_data_summary(result.get("data_summary")),
        "rar_filename": result.get("rar_filename", ""),
        "md5_hash": str(result.get("md5_hash", "")).upper(),
        "file_size": result.get("file_size", ""),
        "disc_number": attach.get("disc_number", ""),
        "burning_date": burning_date,
        "first_evidence_number": evidence_list[0].get("evidence_number", "") if evidence_list else "",
        "created_date": burning_date,
    }
    return flat


# ═══════════════════════════════════════════
# 列表块展开
# ═══════════════════════════════════════════

_LIST_PATTERN = re.compile(r"\{\{#(\w+)\}\}(.*?)\{\{/\1\}\}")


def _expand_all_lists(doc: Document, report: dict):
    """展开文档中所有的 {{#list}}...{{/list}} 块（每个列表只展开一次，删除重复模板段落）"""
    intro = report.get("introduction", {})
    insp = report.get("inspection", {})

    list_data = {
        "evidence_list": intro.get("evidence_list", []),
        "inspectors": intro.get("inspectors", []),
        "process_steps": insp.get("process_steps", []),
    }

    body = doc.element.body
    w_p = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    para_elements = body.findall(w_p)

    # 收集所有带列表标记的段落
    to_expand = []
    for pi, para in enumerate(doc.paragraphs):
        match = _LIST_PATTERN.search(para.text)
        if match:
            to_expand.append((pi, match.group(1), match.group(2)))

    # 标记已处理的列表名，只展开第一个，删除其余
    expanded_names = set()
    elements_to_remove = []

    for pi, list_name, item_template in to_expand:
        if list_name in expanded_names:
            # 重复模板段落，直接删除
            if pi < len(para_elements):
                elements_to_remove.append(para_elements[pi])
        else:
            expanded_names.add(list_name)
            items = list_data.get(list_name, [])
            _expand_list_at_paragraph(doc, para_elements, pi, list_name, item_template, items)

    # 删除重复模板段落
    for el in elements_to_remove:
        try:
            el.getparent().remove(el)
        except Exception:
            pass


def _expand_list_at_paragraph(doc: Document, para_elements: list, pi: int,
                               list_name: str, item_template: str, items: list):
    """在指定段落位置展开列表项"""
    if not items:
        # 空列表：清空段落
        _clear_para(doc.paragraphs[pi])
        return

    original_para = para_elements[pi]
    parent = original_para.getparent()

    # 为每个列表项创建段落副本
    new_elements = []
    for item in items:
        # 深拷贝原段落元素
        new_el = copy.deepcopy(original_para)
        # 在拷贝中替换占位符
        _replace_in_element(new_el, list_name, item_template, item)
        new_elements.append(new_el)

    # 在原始段落之后插入新段落
    insert_after = original_para
    for new_el in new_elements:
        insert_after.addnext(new_el)
        insert_after = new_el

    # 删除原始段落
    parent.remove(original_para)


def _replace_in_element(element, list_name: str, item_template: str, item: dict):
    """在 XML 元素的所有 w:t 节点中替换列表项占位符"""
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    text_elements = element.findall('.//w:t', nsmap)

    if list_name == "evidence_list":
        _fill_evidence_item(text_elements, item_template, item)
    elif list_name == "inspectors":
        _fill_inspector_item(text_elements, item_template, item)
    elif list_name == "process_steps":
        _fill_step_item(text_elements, item_template, item)


def _fill_evidence_item(text_elements, item_template: str, item: dict):
    """填充单个检材条目"""
    # 构建设备描述
    device_type = (
        reviewed_material_display_name(item)
        or item.get("device_name")
        or item.get("model")
        or item.get("device_type", "")
    )
    extractable = item.get("extractable")
    if not isinstance(extractable, bool):
        extractable = any(str(item.get(key, "")).strip() for key in ("imei1", "imei2", "serial_number"))
    display_identifiers = {
        identifier.type: identifier.value
        for identifier in (select_display_identifiers(material_from_legacy_item(item, 0)) if extractable else ())
    }
    imei1 = display_identifiers.get("imei1", "")
    imei2 = display_identifiers.get("imei2", "")
    serial = display_identifiers.get("serial_number", "")
    evidence_number = item.get("evidence_number", "")

    # 构建描述文本（条件字段）
    parts = [f"{device_type}一部"]
    if imei1:
        parts.append(f"IMEI1：{imei1}")
    if imei2:
        parts.append(f"IMEI2：{imei2}")
    if serial:
        parts.append(f"序列号：{serial}")
    full_text = "（" + "；".join(parts[1:]) + "）。" if len(parts) > 1 else "。" if extractable else "（无法提取）。"
    full_text = parts[0] + full_text

    _set_text_elements(text_elements, full_text)


def _fill_inspector_item(text_elements, item_template: str, item: dict):
    """填充单个检查人员条目"""
    name = item.get("name", "")
    unit = item.get("unit", "")
    badge = item.get("badge_number", "")
    full_text = f"{name}，{unit}，警号：{badge}"
    _set_text_elements(text_elements, full_text)


def _fill_step_item(text_elements, item_template: str, item: dict):
    """填充单个检查过程步骤"""
    step_num = item.get("step_number", "")
    content = item.get("content", "")
    full_text = f"{step_num}、{content}"
    _set_text_elements(text_elements, full_text)


def _set_text_elements(text_elements, text: str):
    """设置 XML w:t 元素的文本（保留第一个元素的格式，清空其余）"""
    if text_elements:
        text_elements[0].text = text
        for el in text_elements[1:]:
            el.text = ""


# ═══════════════════════════════════════════
# 简单占位符替换
# ═══════════════════════════════════════════

_SIMPLE_PLACEHOLDER = re.compile(r"\{\{(\w+)\}\}")


def _replace_placeholders(doc: Document, flat: dict):
    """替换文档中所有 {{key}} 占位符（不包括 {{#list}}...{{/list}} 块）"""
    for para in doc.paragraphs:
        for run in para.runs:
            _replace_in_run(run, flat)

    # 表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _replace_in_run(run, flat)


def _replace_in_run(run, flat: dict):
    """在单个 run 中替换所有 {{key}} 占位符"""
    text = run.text
    if not text or "{{" not in text:
        return

    # 移除残留的列表标记
    text = _LIST_PATTERN.sub("", text)

    # 替换简单占位符
    for key, value in flat.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in text:
            text = text.replace(placeholder, str(value))

    run.text = text


def _replace_vml_textbox_placeholders(doc: Document, flat: dict):
    """替换 VML 浮动文本框中的 {{key}} 占位符"""
    w_t = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    w_p = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    body = doc.element.body
    for pict in body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
        txbxs = pict.findall('.//{urn:schemas-microsoft-com:vml}textbox')
        if not txbxs:
            continue
        for para in pict.findall('.//' + w_p):
            all_t = para.findall('.//' + w_t)
            full = "".join(t.text or "" for t in all_t)
            for key, value in flat.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in full:
                    full = full.replace(placeholder, str(value))
            if all_t:
                all_t[0].text = full
                for t in all_t[1:]:
                    t.text = ""


def _replace_header_footer(doc: Document, flat: dict):
    """替换页眉/页脚中的 {{key}} 占位符"""
    for section in doc.sections:
        # 页眉
        for header in (section.header, section.first_page_header, section.even_page_header):
            if header:
                for para in header.paragraphs:
                    for run in para.runs:
                        _replace_in_run(run, flat)
        # 页脚
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            if footer:
                for para in footer.paragraphs:
                    for run in para.runs:
                        _replace_in_run(run, flat)


# ═══════════════════════════════════════════
# 照片处理
# ═══════════════════════════════════════════

def _handle_photos(doc: Document, photo_paths: list[str], report: dict):
    """处理附件2 检材照片 — 标题在图片下方，自适应布局"""
    from lxml import etree

    attachments = report.get("attachments") or {}
    if photo_paths and isinstance(attachments, Mapping) and isinstance(
        attachments.get("photo_groups"), list,
    ):
        _handle_grouped_photos(doc, photo_paths, report)
        return

    v_imagedata = '{urn:schemas-microsoft-com:vml}imagedata'
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body

    # 1. 删除模板图片段落
    to_remove = []
    for pe in body.findall('.//{' + w_ns + '}p'):
        if pe.findall('.//' + v_imagedata):
            to_remove.append(pe)
    for pe in to_remove:
        body.remove(pe)

    # 2. 无照片：清空标题
    if not photo_paths:
        for para in doc.paragraphs:
            if para.text and "检材" in para.text and "照片" in para.text:
                for run in para.runs:
                    run.text = ""
        return

    # 3. 找到照片标题段落（图片将插在标题之前，使标题在图片下方）
    caption_idx = None
    for pi, para in enumerate(doc.paragraphs):
        if para.text and "检材" in para.text and "照片" in para.text:
            caption_idx = pi
            break
    if caption_idx is None:
        return

    caption_elem = doc.paragraphs[caption_idx]._element

    # 4. 根据照片数量选择布局
    if len(photo_paths) == 1:
        img_el = _make_single_image(doc, photo_paths[0], w_ns)
        if img_el is not None:
            caption_elem.addprevious(img_el)
    else:
        tbl = _make_image_grid(doc, photo_paths, w_ns)
        if tbl is not None:
            caption_elem.addprevious(tbl)

    # 5. 更新标题文本中的检材编号
    intro = report.get("introduction", {})
    ev_list = intro.get("evidence_list", [])
    evidence_num = ev_list[0].get("evidence_number", "") if ev_list else ""
    for para in doc.paragraphs:
        if para.text and "检材" in para.text and "照片" in para.text:
            for run in para.runs:
                if "{{first_evidence_number}}" in run.text:
                    run.text = run.text.replace("{{first_evidence_number}}", evidence_num or "")
                elif "first_evidence_number" in run.text:
                    run.text = run.text.replace("first_evidence_number", evidence_num or "")
            break


def _handle_grouped_photos(
    doc: Document, photo_paths: list[str], report: dict,
) -> None:
    """Render report-only exports with the same material-pair layout as formal exports."""
    source_image_ids = photo_values(report)
    groups = material_photo_groups(report)
    pages = build_attachment2_pages(groups)
    assets = validate_attachment2_photos(photo_paths, source_image_ids)
    render_attachment2_pages(
        doc, pages, len(photo_paths), current_template_profile(), assets,
    )


def _make_single_image(doc, photo_path, w_ns):
    """创建居中单张图片段落"""
    from lxml import etree
    if not os.path.isfile(photo_path):
        return None
    try:
        tmp_doc = Document()
        tmp_para = tmp_doc.add_paragraph()
        tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tmp_run = tmp_para.add_run()
        width, height = _fit_image_size(photo_path)
        tmp_run.add_picture(photo_path, width=width, height=height)
        tmp_drawing = tmp_para._element.find('.//{' + w_ns + '}drawing')
        if tmp_drawing is None:
            return None
        _assign_unique_docpr_id(doc, tmp_drawing, _existing_docpr_ids(doc))
        img_part = doc.part.package.get_or_add_image_part(photo_path)
        rid = doc.part.relate_to(img_part, RT.IMAGE)
        for blip in tmp_drawing.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            blip.set(
                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed',
                rid)
        new_p = etree.Element('{' + w_ns + '}p')
        new_pPr = etree.SubElement(new_p, '{' + w_ns + '}pPr')
        etree.SubElement(new_pPr, '{' + w_ns + '}jc').set('{' + w_ns + '}val', 'center')
        new_r = etree.SubElement(new_p, '{' + w_ns + '}r')
        new_r.append(tmp_drawing)
        return new_p
    except Exception:
        return None


def _make_image_grid(doc, photo_paths, w_ns):
    """创建 2 列无边框图片表格（2+ 张照片）"""
    from lxml import etree
    cols = 2
    rows = (len(photo_paths) + cols - 1) // cols
    tbl = etree.Element('{' + w_ns + '}tbl')
    tblPr = etree.SubElement(tbl, '{' + w_ns + '}tblPr')
    tblW = etree.SubElement(tblPr, '{' + w_ns + '}tblW')
    tblW.set('{' + w_ns + '}w', '5000')
    tblW.set('{' + w_ns + '}type', 'pct')
    tblBorders = etree.SubElement(tblPr, '{' + w_ns + '}tblBorders')
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(tblBorders, '{' + w_ns + '}' + bn)
        b.set('{' + w_ns + '}val', 'none')
        b.set('{' + w_ns + '}sz', '0')
        b.set('{' + w_ns + '}space', '0')
        b.set('{' + w_ns + '}color', 'auto')
    tblGrid = etree.SubElement(tbl, '{' + w_ns + '}tblGrid')
    used_docpr_ids = _existing_docpr_ids(doc)
    for c in range(cols):
        gc = etree.SubElement(tblGrid, '{' + w_ns + '}gridCol')
        gc.set('{' + w_ns + '}w', '4500')
    for ri in range(rows):
        tr = etree.SubElement(tbl, '{' + w_ns + '}tr')
        for ci in range(cols):
            pi = ri * cols + ci
            tc = etree.SubElement(tr, '{' + w_ns + '}tc')
            tcPr = etree.SubElement(tc, '{' + w_ns + '}tcPr')
            tcW = etree.SubElement(tcPr, '{' + w_ns + '}tcW')
            tcW.set('{' + w_ns + '}w', '4500')
            tcW.set('{' + w_ns + '}type', 'dxa')
            if pi < len(photo_paths) and os.path.isfile(photo_paths[pi]):
                try:
                    tmp_doc = Document()
                    tmp_para = tmp_doc.add_paragraph()
                    tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tmp_run = tmp_para.add_run()
                    width, height = _fit_image_size(photo_paths[pi])
                    tmp_run.add_picture(photo_paths[pi], width=width, height=height)
                    tmp_drawing = tmp_para._element.find('.//{' + w_ns + '}drawing')
                    if tmp_drawing is not None:
                        _assign_unique_docpr_id(doc, tmp_drawing, used_docpr_ids)
                        img_part = doc.part.package.get_or_add_image_part(photo_paths[pi])
                        rid = doc.part.relate_to(img_part, RT.IMAGE)
                        for blip in tmp_drawing.findall(
                                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                            blip.set(
                                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed',
                                rid)
                        p = etree.SubElement(tc, '{' + w_ns + '}p')
                        pPr = etree.SubElement(p, '{' + w_ns + '}pPr')
                        etree.SubElement(pPr, '{' + w_ns + '}jc').set(
                            '{' + w_ns + '}val', 'center')
                        r = etree.SubElement(p, '{' + w_ns + '}r')
                        r.append(tmp_drawing)
                except Exception:
                    etree.SubElement(tc, '{' + w_ns + '}p')
            else:
                etree.SubElement(tc, '{' + w_ns + '}p')
    return tbl


def _existing_docpr_ids(doc: Document) -> set[int]:
    """读取目标文档已有的 DrawingML 文档属性 ID。"""
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    ids = set()
    for doc_pr in doc.element.findall('.//{' + wp_ns + '}docPr'):
        try:
            ids.add(int(doc_pr.get('id', '0')))
        except ValueError:
            continue
    return ids


def _assign_unique_docpr_id(doc: Document, drawing, used_ids: set[int]):
    """为搬入目标文档的图片分配全局唯一的 wp:docPr/@id。"""
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    doc_pr = drawing.find('.//{' + wp_ns + '}docPr')
    if doc_pr is None:
        return
    next_id = max(used_ids or {0}) + 1
    while next_id in used_ids:
        next_id += 1
    doc_pr.set('id', str(next_id))
    used_ids.add(next_id)


def _fit_image_size(photo_path: str):
    """按原图比例缩放到页面可用区域内，避免图片高度把后续附件推到空白页。"""
    asset = validate_attachment2_photos([photo_path])[0]
    geometry = calculate_fixed_geometry(asset.width_px, asset.height_px)
    return (
        Inches(geometry.render_width_emu / EMU_PER_INCH),
        Inches(geometry.render_height_emu / EMU_PER_INCH),
    )


def _cleanup_attachment_spacing(doc: Document):
    """规范附件间距，并把摘要、签名和日期保持为连续分页块。"""
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body

    def is_empty_paragraph(element):
        if element.tag != '{' + w_ns + '}p':
            return False
        if ''.join(element.itertext()).strip():
            return False
        if element.findall('.//{' + w_ns + '}drawing'):
            return False
        if element.findall('.//{' + w_ns + '}pict'):
            return False
        if any(br.get('{'+w_ns+'}type') == 'page'
               for br in element.findall('.//{' + w_ns + '}br') ):
            return False
        return not element.findall('.//{' + w_ns + '}sectPr')

    def find_paragraph_index(label):
        for index, element in enumerate(list(body)):
            if element.tag == '{' + w_ns + '}p' and label in ''.join(element.itertext()).strip():
                return index
        return None

    attachment_summary_index = find_paragraph_index('1、电子数据提取固定清单')
    if attachment_summary_index is not None:
        children = list(body)
        empty_count = 0
        cursor = attachment_summary_index - 1
        while cursor >= 0 and is_empty_paragraph(children[cursor]):
            empty_count += 1
            cursor -= 1
        while empty_count > 0:
            body.remove(children[attachment_summary_index - 1])
            attachment_summary_index -= 1
            empty_count -= 1
            children = list(body)

        paragraphs = doc.paragraphs
        summary_start = next(
            (index for index, paragraph in enumerate(paragraphs)
             if '1、电子数据提取固定清单' in paragraph.text),
            None,
        )
        attachment1_start = next(
            (index for index, paragraph in enumerate(paragraphs)
             if paragraph.text.strip() == '附件1：'),
            None,
        )
        if (summary_start is not None and attachment1_start is not None
                and summary_start < attachment1_start):
            summary_block = paragraphs[summary_start:attachment1_start]
            first = summary_block[0]
            # Replace the template's three 520-twip blank lines with paragraph
            # spacing. Word keeps it mid-page and suppresses it at an automatic
            # page top, so an independently paginated summary starts normally.
            first.paragraph_format.space_before = Twips(_ATTACHMENT_SUMMARY_GAP_TWIPS)
            for br in list(first._p.findall('.//' + qn('w:br'))):
                if br.get(qn('w:type')) == 'page':
                    br.getparent().remove(br)
            p_pr = first._p.get_or_add_pPr()
            page_break_before = p_pr.find(qn('w:pageBreakBefore'))
            if page_break_before is not None:
                p_pr.remove(page_break_before)
            for index, paragraph in enumerate(summary_block):
                paragraph.paragraph_format.keep_together = True
                if index < len(summary_block) - 1:
                    paragraph.paragraph_format.keep_with_next = True

    attachment2_index = find_paragraph_index('附件2：')
    attachment3_index = find_paragraph_index('附件3：')
    if attachment2_index is None:
        if attachment3_index is None:
            return
        children = list(body)
        while attachment3_index > 0 and is_empty_paragraph(children[attachment3_index - 1]):
            body.remove(children[attachment3_index - 1])
            attachment3_index -= 1
            children = list(body)
        return
    if attachment3_index is None:
        return

    # 表格后的宿主空段落不能留在分页点之前，否则附件2的分页符会制造空白页。
    children = list(body)
    while attachment2_index > 0 and is_empty_paragraph(children[attachment2_index - 1]):
        body.remove(children[attachment2_index - 1])
        attachment2_index -= 1
        children = list(body)

    # 图片、标题和附件3之间不需要模板预留的空行；保留有内容、图片或分页符的节点。
    children = list(body)
    attachment2_index = find_paragraph_index('附件2：')
    attachment3_index = find_paragraph_index('附件3：')
    if attachment2_index is None or attachment3_index is None:
        return
    for element in children[attachment2_index + 1:attachment3_index]:
        if is_empty_paragraph(element):
            body.remove(element)


# ═══════════════════════════════════════════
# 辅助函数
# ═══════════════════════════════════════════

def _clear_para(para):
    """清空段落内容"""
    for run in para.runs:
        run.text = ""


def _apply_required_heading_styles(doc: Document) -> None:
    """Apply stable title formatting after template regions have been rendered."""
    for index, paragraph in enumerate(doc.paragraphs):
        normalized = "".join(paragraph.text.split())
        if index == 0 and normalized:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
        elif normalized == "电子数据提取固定清单":
            for run in paragraph.runs:
                run.bold = True


def _apply_required_attachment_table_styles(doc: Document) -> None:
    """Apply stable Attachment 1 header and Latin wrapping rules."""
    if not doc.tables:
        return
    attachment_table = doc.tables[0]
    for cell in attachment_table.rows[0].cells:
        if "".join(cell.text.split()) not in {"电子数据", "来源"}:
            continue
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph_pr = paragraph._p.get_or_add_pPr()
            indent = paragraph_pr.find(qn("w:ind"))
            if indent is not None:
                paragraph_pr.remove(indent)
            for run_properties in paragraph._p.findall(".//" + qn("w:rPr")):
                for local_name in ("spacing", "w", "fitText"):
                    node = run_properties.find(qn(f"w:{local_name}"))
                    if node is not None:
                        run_properties.remove(node)
    for row in attachment_table.rows:
        for cell_index in (1, 4):
            if cell_index < len(row.cells):
                allow_latin_character_wrap(row.cells[cell_index]._tc)


def _remove_comments(doc: Document):
    """删除文档中所有批注引用元素"""
    body = doc.element.body
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for tag in ['commentRangeStart', 'commentRangeEnd', 'commentReference']:
        for el in body.findall('.//{' + w_ns + '}' + tag):
            el.getparent().remove(el)


# ═══════════════════════════════════════════
# 表格提取清单展开
# ═══════════════════════════════════════════

def _expand_extract_table(doc: Document, report: dict):
    """展开附件1 电子数据提取固定清单表格"""
    extract_list = report.get("attachments", {}).get("extract_list", {})
    rows = extract_list.get("rows", [])

    if not doc.tables:
        return

    table = doc.tables[0]
    if len(table.rows) < 2:
        return

    signing_row_idx = len(table.rows) - 1
    template_row = table.rows[1]

    if not rows:
        # 空数据：清除所有数据行文本，在 Row 1 画左下→右上对角线
        for ri in range(1, signing_row_idx):
            if ri < len(table.rows):
                _clear_row_and_draw_diagonal(table.rows[ri], is_first_data_row=(ri == 1))
        return

    # 用实际数据填充模板行
    first_row = rows[0]
    _fill_table_row(template_row, first_row)

    # 为额外数据行插入新行
    tbl = table._tbl
    last_data_row_el = template_row._tr

    for extra_row in rows[1:]:
        new_tr = copy.deepcopy(template_row._tr)
        last_data_row_el.addnext(new_tr)
        last_data_row_el = new_tr
        from docx.table import _Row
        new_row_obj = _Row(new_tr, table)
        _fill_table_row(new_row_obj, extra_row)

    # 清除多余的空数据行
    for ri in range(1 + len(rows), signing_row_idx):
        if ri < len(table.rows):
            for cell in table.rows[ri].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ""


def _clear_row_and_draw_diagonal(row, is_first_data_row: bool = True):
    """清除行内所有文本，并在第一个数据行画对角线（左下→右上）"""
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    }
    for ci, cell in enumerate(row.cells):
        # 清除文本
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""
        # 在第一行画对角线
        if is_first_data_row:
            tc = cell._tc
            tcPr = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            if tcPr is None:
                from lxml import etree
                tcPr = etree.SubElement(tc, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            # 添加 tcBorders
            tcBorders = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
            if tcBorders is None:
                from lxml import etree
                tcBorders = etree.SubElement(tcPr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
            # 添加 tr2bl（左下→右上），即 top-right to bottom-left 对角线
            tr2bl = etree.SubElement(tcBorders, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr2bl')
            tr2bl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
            tr2bl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '4')
            tr2bl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '0')
            tr2bl.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '000000')


def _fill_table_row(row, item: dict):
    """填充表格行数据"""
    cell_keys = ["no", "electronic_data", "source", "extraction_method", "md5_hash"]
    for ci, key in enumerate(cell_keys):
        if ci < len(row.cells):
            value = item.get(key, "")
            if key == "md5_hash":
                value = str(value).upper()
            elif key == "source":
                source = str(value).strip()
                if source.endswith("内提取") and not source.endswith("检材内提取"):
                    source = source[:-3] + "检材内提取"
                value = "\n".join(attachment1_source_lines(source))
            cell = row.cells[ci]
            if key == "source":
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for para in cell.paragraphs:
                for run in para.runs:
                    # 替换占位符
                    placeholder = f"{{{{extract_{key}}}}}"
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
                    elif run.text == placeholder:
                        run.text = str(value)
