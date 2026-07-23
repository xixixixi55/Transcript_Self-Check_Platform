"""OOXML helpers for Attachment2 image relationships and IDs."""

from __future__ import annotations

from typing import Any

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Emu
from lxml import etree

from .attachment2_image_service import Attachment2PhotoAsset, calculate_fixed_geometry
from .attachment_plan_service import AttachmentPlanError
from .docx_attachment_xml_service import W_NS, qn, text_of
from .template_profile_service import CurrentTemplateProfile

_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def build_attachment2_drawing(
    doc: Any, asset: Attachment2PhotoAsset, profile: CurrentTemplateProfile,
    used_drawing_ids: set[int],
) -> Any:
    geometry = calculate_fixed_geometry(
        asset.width_px, asset.height_px,
        slot_width_emu=profile.attachment2_slot_width_emu,
        slot_height_emu=profile.attachment2_slot_height_emu,
    )
    temporary = Document()
    run = temporary.add_paragraph().add_run()
    run.add_picture(
        asset.path, width=Emu(geometry.render_width_emu),
        height=Emu(geometry.render_height_emu),
    )
    drawing = run._r.find(".//%s" % qn(_WP_NS, "inline"))
    if drawing is None:
        raise AttachmentPlanError("ATTACHMENT_PLAN_INVALID", "附件2图片绘图结构无效。")
    image_part = doc.part.package.get_or_add_image_part(asset.path)
    relationship_id = doc.part.relate_to(image_part, RT.IMAGE)
    for blip in drawing.findall(".//%s" % qn(_A_NS, "blip")):
        blip.set(qn(_REL_NS, "embed"), relationship_id)
    assign_drawing_ids(drawing, used_drawing_ids)
    wrapper = etree.Element(qn(W_NS, "drawing"))
    wrapper.append(drawing)
    return wrapper


def existing_drawing_ids(root: Any) -> set[int]:
    values: set[int] = set()
    for element in root.iter():
        if element.tag in {qn(_WP_NS, "docPr"), qn(_PIC_NS, "cNvPr")}:
            try:
                values.add(int(element.get("id", "0")))
            except ValueError:
                continue
    return values


def assign_drawing_ids(root: Any, used: set[int]) -> None:
    next_id = max(used or {0}) + 1
    for element in root.iter():
        if element.tag not in {qn(_WP_NS, "docPr"), qn(_PIC_NS, "cNvPr")}:
            continue
        while next_id in used:
            next_id += 1
        element.set("id", str(next_id))
        used.add(next_id)
        next_id += 1


def find_attachment2_paragraph(body: Any, anchor: str, exact: bool = False) -> Any | None:
    for element in body.findall("./%s" % qn(W_NS, "p")):
        value = text_of(element)
        if (value == anchor if exact else anchor in value):
            return element
    return None


def twips(emu: int) -> int:
    return round(emu / 635)


__all__ = [
    "assign_drawing_ids", "build_attachment2_drawing", "existing_drawing_ids",
    "find_attachment2_paragraph", "twips",
]
