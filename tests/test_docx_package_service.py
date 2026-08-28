"""版本化 DOCX 包指纹与 ZIP 安全门控的属性测试。"""

import os
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "packages", "backend"))

from app.services.docx_package_service import (  # noqa: E402
    DocxPackageError,
    OOXML_PACKAGE_FINGERPRINT_ALGORITHM,
    _validate_entry_name,
    compute_ooxml_package_fingerprint,
)
from app.services.template_profile_service import (  # noqa: E402
    CURRENT_TEMPLATE_PACKAGE_FINGERPRINT,
    TemplateProfileError,
    current_template_profile,
    validate_current_template_profile,
    validate_template_package_fingerprint,
)
from docx import Document  # noqa: E402


ROOT = Path(__file__).parents[1]
TEMPLATE = ROOT / "word_templates" / "template.docx"
REFERENCE = ROOT / "2026报告模板（one压缩包）最终提交.docx"


def _write_package(path: Path, entries: list[tuple[str, bytes]], *, level: int = 6) -> None:
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=level) as package:
        for name, content in entries:
            package.writestr(name, content)


def _entries() -> list[tuple[str, bytes]]:
    return [
        ("word/document.xml", b"<w:document/>"),
        ("word/footer1.xml", b"footer"),
        ("word/vml.xml", b"<v:shape id='original'/>") ,
        ("docProps/core.xml", b"<core>original</core>"),
        ("word/comments.xml", b"<comments>original</comments>"),
    ]


def test_container_metadata_does_not_change_fingerprint(tmp_path):
    first, second, third = (tmp_path / name for name in ("a.docx", "b.docx", "c.docx"))
    _write_package(first, _entries(), level=1)
    _write_package(second, list(reversed(_entries())), level=9)
    with zipfile.ZipFile(third, "w") as package:
        for name, content in _entries():
            info = zipfile.ZipInfo(name, date_time=(2030, 1, 2, 3, 4, 5))
            info.compress_type = zipfile.ZIP_DEFLATED
            package.writestr(info, content)
    assert compute_ooxml_package_fingerprint(first) == compute_ooxml_package_fingerprint(second)
    assert compute_ooxml_package_fingerprint(first) == compute_ooxml_package_fingerprint(third)


@pytest.mark.parametrize("entry", [
    ("word/document.xml", b"<w:document/>changed"),
    ("word/footer1.xml", b"footer changed"),
    ("word/vml.xml", b"<v:shape id='changed'/>") ,
    ("docProps/core.xml", b"<core>changed</core>"),
    ("word/comments.xml", b"<comments>changed</comments>"),
])
def test_package_content_changes_fingerprint(tmp_path, entry):
    original = tmp_path / "original.docx"
    changed = tmp_path / "changed.docx"
    _write_package(original, _entries())
    name, content = entry
    changed_entries = [(name if current == name else current, content if current == name else value)
                       for current, value in _entries()]
    if name not in {current for current, _ in _entries()}:
        changed_entries.append(entry)
    _write_package(changed, changed_entries)
    assert compute_ooxml_package_fingerprint(original) != compute_ooxml_package_fingerprint(changed)


@pytest.mark.parametrize("entries", [
    _entries() + [("word/new.xml", b"new")],
    [("word/document.xml", b"<w:document/>")],
    [("word/renamed.xml", b"<w:document/>")],
])
def test_entry_set_changes_fingerprint(tmp_path, entries):
    original = tmp_path / "original.docx"
    changed = tmp_path / "changed.docx"
    _write_package(original, _entries())
    _write_package(changed, entries)
    assert compute_ooxml_package_fingerprint(original) != compute_ooxml_package_fingerprint(changed)


def test_current_template_matches_registered_fingerprint_and_profile():
    current_fingerprint = compute_ooxml_package_fingerprint(TEMPLATE)

    assert current_fingerprint == CURRENT_TEMPLATE_PACKAGE_FINGERPRINT
    profile = current_template_profile()
    assert profile.fingerprint_algorithm == OOXML_PACKAGE_FINGERPRINT_ALGORITHM
    assert validate_current_template_profile(
        str(TEMPLATE), Document(str(TEMPLATE)),
    ).package_fingerprint == profile.package_fingerprint


def test_profile_rejects_any_unbalanced_horizontal_body_indent():
    doc = Document(str(TEMPLATE))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    indent = doc.element.body.find(f"./{namespace}p/{namespace}pPr/{namespace}ind")
    assert indent is not None
    indent.set(f"{namespace}left", "500")
    indent.set(f"{namespace}right", "0")

    with pytest.raises(TemplateProfileError, match="未居中"):
        validate_current_template_profile(str(TEMPLATE), doc)


def test_profile_requires_balanced_horizontal_body_indents_to_be_present():
    doc = Document(str(TEMPLATE))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for indent in doc.element.body.findall(
        f"./{namespace}p/{namespace}pPr/{namespace}ind",
    ):
        for name in ("left", "right", "leftChars", "rightChars"):
            indent.attrib.pop(f"{namespace}{name}", None)

    with pytest.raises(TemplateProfileError, match="未居中"):
        validate_current_template_profile(str(TEMPLATE), doc)


def test_profile_rejects_visible_title_tabs_and_inset_structural_headings():
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with_title_tab = Document(str(TEMPLATE))
    title_run = with_title_tab.element.body.find(f"./{namespace}p/{namespace}r")
    assert title_run is not None
    title_run.append(with_title_tab.element.body.makeelement(f"{namespace}tab"))
    with pytest.raises(TemplateProfileError, match="标题、层级或横线未居中"):
        validate_current_template_profile(str(TEMPLATE), with_title_tab)

    inset_heading = Document(str(TEMPLATE))
    heading = next(
        paragraph for paragraph in inset_heading.paragraphs
        if paragraph.text.strip() == "一、绪论"
    )
    indent = heading._p.find(f"./{namespace}pPr/{namespace}ind")
    assert indent is not None
    indent.set(f"{namespace}firstLine", "614")
    with pytest.raises(TemplateProfileError, match="标题、层级或横线未居中"):
        validate_current_template_profile(str(TEMPLATE), inset_heading)


def test_profile_rejects_off_center_fixed_horizontal_rule(tmp_path):
    with zipfile.ZipFile(TEMPLATE) as package:
        entries = [(name, package.read(name)) for name in package.namelist()]
    footer_name = "word/footer1.xml"
    footer = ET.fromstring(dict(entries)[footer_name])
    line = footer.find(".//{urn:schemas-microsoft-com:vml}line")
    assert line is not None
    line.set("from", "0pt,4pt")
    changed_entries = [
        (
            name,
            ET.tostring(footer, encoding="utf-8", xml_declaration=True)
            if name == footer_name else content,
        )
        for name, content in entries
    ]
    invalid = tmp_path / "SYNTHETIC-off-center-rule.docx"
    _write_package(invalid, changed_entries)
    fingerprint = compute_ooxml_package_fingerprint(invalid)

    with pytest.raises(TemplateProfileError, match="标题、层级或横线未居中"):
        validate_current_template_profile(
            str(invalid), Document(str(invalid)), fingerprint,
        )


def test_profile_requires_fixed_title_and_document_number_slots():
    without_title = Document(str(TEMPLATE))
    for node in without_title.element.body[0].xpath(".//w:t"):
        node.text = ""
    with pytest.raises(TemplateProfileError, match="固定标题槽"):
        validate_current_template_profile(str(TEMPLATE), without_title)

    moved_title = Document(str(TEMPLATE))
    first, second = moved_title.element.body[0], moved_title.element.body[1]
    title_text = "".join(first.xpath(".//w:t/text()"))
    for node in first.xpath(".//w:t"):
        node.text = ""
    second_text_nodes = second.xpath(".//w:t")
    assert second_text_nodes
    second_text_nodes[0].text = title_text
    for node in second_text_nodes[1:]:
        node.text = ""
    with pytest.raises(TemplateProfileError, match="固定标题槽"):
        validate_current_template_profile(str(TEMPLATE), moved_title)


def test_accepted_reference_does_not_match_current_profile():
    if not REFERENCE.is_file():
        pytest.skip("甲方参考模板不在当前工作区")
    assert compute_ooxml_package_fingerprint(REFERENCE) != current_template_profile().package_fingerprint


@pytest.mark.parametrize("name", [
    "/word/document.xml",
    "word/../document.xml",
    "word//document.xml",
    "C:/document.xml",
    "//?/C:/document.xml",
])
def test_unsafe_entry_names_are_rejected(tmp_path, name):
    package = tmp_path / "unsafe.docx"
    _write_package(package, [(name, b"x")])
    with pytest.raises(DocxPackageError):
        compute_ooxml_package_fingerprint(package)


def test_empty_entry_name_is_rejected():
    with pytest.raises(DocxPackageError):
        _validate_entry_name("", set(), set())


def test_backslash_entry_name_is_rejected():
    with pytest.raises(DocxPackageError):
        _validate_entry_name(r"word\document.xml", set(), set())


@pytest.mark.filterwarnings("ignore:Duplicate name:UserWarning")
def test_duplicate_and_case_folded_entries_are_rejected(tmp_path):
    package = tmp_path / "duplicate.docx"
    with zipfile.ZipFile(package, "w") as archive:
        archive.writestr("word/document.xml", b"one")
        archive.writestr("word/document.xml", b"two")
    with pytest.raises(DocxPackageError):
        compute_ooxml_package_fingerprint(package)

    folded = tmp_path / "folded.docx"
    _write_package(folded, [("word/document.xml", b"one"), ("word/DOCUMENT.XML", b"two")])
    with pytest.raises(DocxPackageError):
        compute_ooxml_package_fingerprint(folded)


def test_non_zip_and_encrypted_entries_are_rejected(tmp_path):
    non_zip = tmp_path / "not-a-docx.docx"
    non_zip.write_bytes(b"not a ZIP")
    with pytest.raises(DocxPackageError):
        compute_ooxml_package_fingerprint(non_zip)

    encrypted = tmp_path / "encrypted.docx"
    _write_package(encrypted, [("word/document.xml", b"x")])
    data = bytearray(encrypted.read_bytes())
    local = data.index(b"PK\x03\x04")
    central = data.index(b"PK\x01\x02")
    data[local + 6:local + 8] = (1).to_bytes(2, "little")
    data[central + 8:central + 10] = (1).to_bytes(2, "little")
    encrypted.write_bytes(data)
    with pytest.raises(DocxPackageError):
        compute_ooxml_package_fingerprint(encrypted)


def test_profile_rejects_invalid_package_with_stable_error(tmp_path):
    invalid = tmp_path / "invalid-template.docx"
    invalid.write_bytes(b"not a ZIP")
    with pytest.raises(TemplateProfileError) as error:
        validate_template_package_fingerprint(str(invalid))
    assert error.value.code == "TEMPLATE_PACKAGE_INVALID"
