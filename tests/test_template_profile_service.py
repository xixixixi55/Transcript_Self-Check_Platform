"""T018 registered-template fingerprint, rule, and structure validation tests."""

from __future__ import annotations

import json
import os
import shutil
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "packages", "backend"))

from app.repository.template_approval_repository import TemplateApprovalRepository  # noqa: E402
from app.repository.template_registry_repository import TemplateRegistryRepository  # noqa: E402
from app.repository.workbench_database import WorkbenchDatabase  # noqa: E402
from app.services.docx_package_service import compute_ooxml_package_fingerprint  # noqa: E402
from app.services.template_profile_service import (  # noqa: E402
    CURRENT_TEMPLATE_PACKAGE_FINGERPRINT,
    CURRENT_TEMPLATE_VALIDATION_RULE,
    validate_registered_template,
)
from app.services.template_registry_service import TemplateRegistryService  # noqa: E402
from test_legacy_report_projection_service import _report  # noqa: E402

ROOT = Path(__file__).parents[1]
SOURCE_TEMPLATE = ROOT / "word_templates" / "template.docx"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS = "urn:schemas-microsoft-com:vml"


def setup_version(tmp_path: Path, template: Path, *, rules=None, status="approved"):
    database = WorkbenchDatabase(tmp_path / "workbench.sqlite3", "SYNTHETIC-PROFILE")
    registry = TemplateRegistryRepository(database, (tmp_path,))
    approvals = TemplateApprovalRepository(database, registry)
    reference = {"template_id": "template-SYNTHETIC-profile", "version": "1.0.0"}
    registry.register({
        "schema_version": 1, "template_ref": reference,
        "display_name": "SYNTHETIC profile template",
        "fingerprint": compute_ooxml_package_fingerprint(template),
        "validation_rules": rules or [CURRENT_TEMPLATE_VALIDATION_RULE],
        "asset_id": "asset-SYNTHETIC-profile",
        "registered_at": "2026-07-30T00:00:00+00:00",
    }, template)
    approvals.record(reference, {
        "approval_record_id": "approval-SYNTHETIC-profile",
        "status": status, "acceptance_summary": "SYNTHETIC acceptance",
        "recorded_at": "2026-07-30T01:00:00+00:00",
    })
    return registry, approvals, reference


def test_approved_registered_template_validates_with_safe_public_result(tmp_path: Path):
    template = tmp_path / "SYNTHETIC-valid.docx"
    shutil.copy2(SOURCE_TEMPLATE, template)
    registry, approvals, reference = setup_version(tmp_path, template)

    result = validate_registered_template(registry, approvals, reference)

    assert result["valid"] is True
    assert result["template"]["fingerprint"] == CURRENT_TEMPLATE_PACKAGE_FINGERPRINT
    assert "internal_locator" not in result["template"]
    assert str(tmp_path) not in repr(result)


def test_fingerprint_change_is_rejected_before_generation(tmp_path: Path):
    template = tmp_path / "SYNTHETIC-drift.docx"
    shutil.copy2(SOURCE_TEMPLATE, template)
    registry, approvals, reference = setup_version(tmp_path, template)
    with zipfile.ZipFile(template, "a") as package:
        package.writestr("word/SYNTHETIC-drift.xml", b"<SYNTHETIC/>")

    result = validate_registered_template(registry, approvals, reference)

    assert result == {
        "valid": False,
        "error_code": "TEMPLATE_FINGERPRINT_MISMATCH",
        "safe_summary": "所选模板指纹校验失败。",
    }
    assert TemplateRegistryService(
        registry.database, registry, approvals,
    ).list_available() == []


def test_unknown_rule_and_unapproved_version_use_stable_safe_codes(tmp_path: Path):
    template = tmp_path / "SYNTHETIC-rule.docx"
    shutil.copy2(SOURCE_TEMPLATE, template)
    registry, approvals, reference = setup_version(
        tmp_path, template,
        rules=[{"rule_id": "rule-SYNTHETIC-unknown", "version": "1.0.0"}],
        status="pending",
    )
    result = validate_registered_template(registry, approvals, reference)
    assert result["error_code"] == "TEMPLATE_NOT_APPROVED"
    assert str(template) not in repr(result)

    approved_root = tmp_path / "approved"
    approved_root.mkdir()
    approved_template = approved_root / "SYNTHETIC-unknown-rule.docx"
    shutil.copy2(SOURCE_TEMPLATE, approved_template)
    registry, approvals, reference = setup_version(
        approved_root, approved_template,
        rules=[{"rule_id": "rule-SYNTHETIC-unknown", "version": "1.0.0"}],
    )
    result = validate_registered_template(registry, approvals, reference)
    assert result["error_code"] == "TEMPLATE_RULE_VALIDATION_FAILED"


def test_case_selection_only_updates_version_ref_and_invalidates_word(tmp_path: Path):
    template = tmp_path / "SYNTHETIC-case-selection.docx"
    shutil.copy2(SOURCE_TEMPLATE, template)
    registry, approvals, reference = setup_version(tmp_path, template)
    database = registry.database
    report = _report()
    report["inspection"].pop("primary_software", None)
    with database.transaction() as connection:
        connection.execute(
            "INSERT INTO case_shells VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                "case-SYNTHETIC-template", 1, "SYNTHETIC-001", "SYNTHETIC case",
                "SYNTHETIC summary", "source-SYNTHETIC-template",
                "task-SYNTHETIC-template", "review_ready", 1, 4,
                "2026-07-30T00:00:00+00:00", "2026-07-30T00:00:00+00:00",
            ),
        )
        connection.execute(
            "INSERT INTO case_drafts VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (
                "case-SYNTHETIC-template", 1, json.dumps(report),
                "legacy-v1", "{}", "[]", None, "plan-SYNTHETIC-existing",
                "review_ready", 4, "2026-07-30T00:00:00+00:00",
                "2026-07-30T00:00:00+00:00",
            ),
        )
    service = TemplateRegistryService(database, registry, approvals)
    assert service.references.drafts.get("case-SYNTHETIC-template")["template_ref"] is None

    result = service.select_for_case(
        "case-SYNTHETIC-template", reference, expected_revision=4,
    )

    assert result["draft"]["template_ref"] == reference
    assert result["draft"]["archive_plan_id"] == "plan-SYNTHETIC-existing"
    assert result["draft"]["revision"] == 5
    assert result["impact"] == {
        "word_artifact_validity": "invalidated_by_template_change",
        "archive_plan_changed": False, "archive_task_created": False,
        "manifest_changed": False, "disc_mapping_changed": False,
    }
    with database.connect() as connection:
        assert connection.execute("SELECT COUNT(*) FROM task_records").fetchone()[0] == 0
        assert connection.execute("SELECT COUNT(*) FROM archive_attempts").fetchone()[0] == 0
    approvals.record(reference, {
        "approval_record_id": "approval-SYNTHETIC-revoked",
        "status": "revoked", "acceptance_summary": "SYNTHETIC revoked",
        "recorded_at": "2026-07-30T02:00:00+00:00",
    })
    assert service.references.drafts.get("case-SYNTHETIC-template")["template_ref"] == reference


@pytest.mark.parametrize("mutation", ["vml", "pagination", "table", "attachment"])
def test_profile_rule_rejects_vml_pagination_table_and_attachment_drift(
    tmp_path: Path, mutation: str,
):
    template = tmp_path / f"SYNTHETIC-{mutation}.docx"
    document = Document(str(SOURCE_TEMPLATE))
    body = document.element.body
    if mutation == "vml":
        for node in body.findall(f".//{{{V_NS}}}textbox"):
            node.getparent().remove(node)
    elif mutation == "pagination":
        paragraph = next(item for item in document.paragraphs if item.text.strip() == "附件2：")
        for node in paragraph._element.findall(f".//{{{W_NS}}}br"):
            node.getparent().remove(node)
    elif mutation == "table":
        body.remove(document.tables[0]._element)
    else:
        paragraph = next(
            item for item in document.paragraphs
            if item.text.strip() == "检材{{first_evidence_number}}照片"
        )
        body.remove(paragraph._element)
    document.save(template)
    registry, approvals, reference = setup_version(tmp_path, template)

    result = validate_registered_template(registry, approvals, reference)

    assert result["valid"] is False
    assert result["error_code"] == "TEMPLATE_RULE_VALIDATION_FAILED"
    assert str(template) not in repr(result)
