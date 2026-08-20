"""模板填充回归测试：VML、默认数据摘要和附件分页。"""

import os
import struct
import sys
import zipfile
import zlib
from pathlib import Path
import xml.etree.ElementTree as ET

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "packages", "backend"))

from app.services.template_filler_service import _flatten_report, fill_template
from app.services.docx_package_service import compute_ooxml_package_fingerprint
from app.services.attachment2_image_service import (
    ATTACHMENT2_DUAL_GROUP_SLOT_HEIGHT_EMU,
    ATTACHMENT2_SLOT_HEIGHT_EMU,
    calculate_fixed_geometry,
)
from app.services.template_profile_service import current_template_profile
from synthetic_report_builders import build_archive_manifest


_ROOT = Path(__file__).parents[1]
_TEMPLATE = _ROOT / "word_templates" / "template.docx"
_DEFAULT_SUMMARY = "即时通讯、手机信息"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _report(data_summary_marker=...):
    result = {
        "evidence_number": "JC01",
        "software_name": "测试工具",
        "software_version": "1.0",
        "rar_filename": "case.rar",
        "md5_hash": "a" * 32,
        "file_size": "123",
    }
    if data_summary_marker is not ...:
        result["data_summary"] = data_summary_marker
    return {
        "title": "电子数据检查笔录",
        "document_number": "SYN-TEST〔2026〕000号",
        "introduction": {
            "entrust_unit": "测试单位",
            "entrust_persons": ["测试人员"],
            "entrust_time": "2026年7月16日",
            "case_summary": "测试案件",
            "evidence_list": [{"evidence_number": "JC01", "device_type": "测试手机"}],
            "inspection_requirement": "测试要求",
            "inspection_time_range": "2026年7月16日",
            "inspectors": [],
            "inspection_place": "测试鉴定中心",
        },
        "inspection": {
            "method": "测试方法",
            "hardware_device": "测试设备",
            "software_tools": [],
            "process_steps": [],
            "result": result,
        },
        "attachments": {
            "extract_list": {"rows": []},
            "photo_ids": [],
            "disc_number": "TEST-DISC",
            "burning_date": "2026年7月16日",
        },
    }


@pytest.mark.parametrize("value", [None, "", "   "])
def test_data_summary_blank_values_use_fixed_default(value):
    report = _report(value)
    assert _flatten_report(report)["data_summary"] == _DEFAULT_SUMMARY


def test_data_summary_missing_uses_fixed_default():
    assert _flatten_report(_report())["data_summary"] == _DEFAULT_SUMMARY


def test_data_summary_normal_value_is_preserved():
    assert _flatten_report(_report("  通讯录  "))["data_summary"] == "通讯录"


def test_flatten_report_names_all_evidence_items_in_result():
    report = _report()
    report["introduction"]["evidence_list"].append({
        "evidence_number": "JC02",
        "device_type": "测试平板",
    })
    assert _flatten_report(report)["evidence_number"] == "JC01、JC02"


def test_flatten_report_normalizes_entrust_person_separators():
    report = _report()
    report["introduction"]["entrust_persons"] = ["SYNTHETIC-A；SYNTHETIC-B", "SYNTHETIC-C"]

    assert _flatten_report(report)["entrust_persons_text"] == (
        "SYNTHETIC-A、SYNTHETIC-B、SYNTHETIC-C"
    )


def test_flatten_report_uses_burning_date_for_attachment_summary_signature():
    flat = _flatten_report(_report())

    assert flat["created_date"] == "2026年7月16日"


def test_flatten_report_combines_entrust_unit_prefix_without_separator():
    report = _report()
    report["introduction"]["entrust_unit_prefix"] = " SYNTHETIC-公安分局 "
    report["introduction"]["entrust_unit"] = " SYNTHETIC-派出所 "

    assert _flatten_report(report)["entrust_unit"] == "SYNTHETIC-公安分局SYNTHETIC-派出所"

    report["introduction"]["entrust_unit_prefix"] = ""
    assert _flatten_report(report)["entrust_unit"] == "SYNTHETIC-派出所"


def test_word_titles_md5_and_legacy_extract_source_are_normalized(tmp_path):
    report = _report()
    report["introduction"]["entrust_unit_prefix"] = "SYNTHETIC-公安分局"
    report["introduction"]["entrust_unit"] = "SYNTHETIC-派出所"
    report["attachments"]["extract_list"]["rows"] = [{
        "no": "1",
        "electronic_data": "SYNTHETIC.rar",
        "source": "JC01内提取",
        "extraction_method": "SYNTHETIC/TEST",
        "md5_hash": "abcdef0123456789abcdef0123456789",
    }]
    output = tmp_path / "normalized-format.docx"
    fill_template(report, str(_TEMPLATE), str(output))

    document = Document(output)
    title = next(p for p in document.paragraphs if p.text.strip() == "电子数据检查笔录")
    extract_heading = next(
        p for p in document.paragraphs if p.text.strip() == "电子数据提取固定清单"
    )
    assert title.alignment == 1
    assert title.runs and all(run.bold for run in title.runs if run.text)
    assert extract_heading.runs and all(run.bold for run in extract_heading.runs if run.text)
    header_cells = document.tables[0].rows[0].cells
    for cell_index, expected_text in ((1, "电子数据"), (2, "来源")):
        paragraph = header_cells[cell_index].paragraphs[0]
        assert paragraph.text == expected_text
        assert paragraph.alignment == 1
        assert paragraph._p.pPr.find(qn("w:ind")) is None
        assert not paragraph._p.findall(".//" + qn("w:rPr") + "/" + qn("w:spacing"))
    assert any(
        "委托单位：SYNTHETIC-公安分局SYNTHETIC-派出所" in paragraph.text
        for paragraph in document.paragraphs
    )
    assert document.tables[0].rows[1].cells[2].text.strip().splitlines() == [
        "JC01", "检材内提取",
    ]
    assert all(
        paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        for paragraph in document.tables[0].rows[1].cells[2].paragraphs
    )
    assert document.tables[0].rows[1].cells[4].text.strip() == "ABCDEF0123456789ABCDEF0123456789"
    for cell_index in (1, 4):
        for paragraph in document.tables[0].rows[1].cells[cell_index].paragraphs:
            word_wrap = paragraph._p.pPr.find(qn("w:wordWrap"))
            assert word_wrap is not None
            assert word_wrap.get(qn("w:val")) == "off"


def test_manifest_render_removes_spread_formatting_from_short_extract_headers(tmp_path):
    report = _report()
    report["inspection"]["primary_software"] = {
        "name": "SYNTHETIC 测试取证软件",
        "version": "1.0",
        "confirmation_status": "confirmed_by_report",
    }
    report["inspection"]["software_tools"] = [
        {"name": "WinRAR压缩管理软件", "version": "6.24"},
        {"name": "Python hashlib", "version": "3.12"},
    ]
    template = Document(_TEMPLATE)
    for cell_index in (1, 2):
        paragraph = template.tables[0].rows[0].cells[cell_index].paragraphs[0]
        paragraph_pr = paragraph._p.get_or_add_pPr()
        indent = paragraph_pr.find(qn("w:ind"))
        if indent is None:
            indent = OxmlElement("w:ind")
            paragraph_pr.append(indent)
        indent.set(qn("w:right"), "720")
        for run in paragraph.runs:
            run_properties = run._r.get_or_add_rPr()
            for local_name, value in (("spacing", "120"), ("w", "150"), ("fitText", "600")):
                node = OxmlElement(f"w:{local_name}")
                node.set(qn("w:val"), value)
                run_properties.append(node)
    modified_template = tmp_path / "spread-header-template.docx"
    template.save(modified_template)
    output = tmp_path / "normalized-manifest-format.docx"

    fill_template(
        report, str(modified_template), str(output), [], build_archive_manifest(1),
        compute_ooxml_package_fingerprint(modified_template),
    )

    document = Document(output)
    for cell_index, expected_text in ((1, "电子数据"), (2, "来源")):
        paragraph = document.tables[0].rows[0].cells[cell_index].paragraphs[0]
        assert paragraph.text == expected_text
        assert paragraph.alignment == 1
        assert paragraph._p.pPr.find(qn("w:ind")) is None
        for local_name in ("spacing", "w", "fitText"):
            assert not paragraph._p.findall(
                ".//" + qn("w:rPr") + "/" + qn(f"w:{local_name}")
            )


def test_manifest_disc_date_overrides_attachment_summary_signature_date(tmp_path):
    report = _report()
    report["inspection"]["primary_software"] = {
        "name": "测试取证软件",
        "version": "1.0",
        "confirmation_status": "confirmed_by_report",
    }
    report["inspection"]["software_tools"] = [
        {"name": "WinRAR压缩管理软件", "version": "6.24"},
        {"name": "Python hashlib", "version": "3.12"},
    ]
    output = tmp_path / "manifest-signature-date.docx"
    fill_template(report, str(_TEMPLATE), str(output), [], build_archive_manifest())

    paragraphs = Document(output).paragraphs
    signature_index = next(
        index for index, paragraph in enumerate(paragraphs)
        if "检查人签名" in paragraph.text
    )
    signature_date = next(
        paragraph.text.strip()
        for paragraph in paragraphs[signature_index + 1:signature_index + 5]
        if paragraph.text.strip()
    )

    assert signature_date == "2026年7月6日"


def test_fill_template_combines_all_evidence_numbers_in_result_sentence(tmp_path):
    report = _report()
    report["introduction"]["evidence_list"].append({
        "evidence_number": "JC02",
        "device_type": "测试平板",
    })
    output = tmp_path / "combined-result.docx"
    fill_template(report, str(_TEMPLATE), str(output))

    with zipfile.ZipFile(output) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")

    assert "经对编号为JC01、JC02号检材使用测试工具" in document_xml
    assert "；经对编号为JC02号检材" not in document_xml


def test_evidence_renderer_preserves_material_type_and_identifier_contracts(tmp_path):
    report = _report()
    report["introduction"]["evidence_list"] = [
        {
            "id": "phone", "evidence_number": "JC-PHONE", "device_type": "手机",
            "material_type": "phone", "material_type_status": "confirmed_by_user",
            "material_type_source": "user", "imei1": "111111111111111",
            "serial_number": "PHONE-SERIAL-MUST-NOT-RENDER",
        },
        {
            "id": "tablet", "evidence_number": "JC-TABLET", "device_type": "平板",
            "material_type": "tablet", "material_type_status": "confirmed_by_user",
            "material_type_source": "user", "imei1": "222222222222222",
            "serial_number": "TABLET-SERIAL-MUST-RENDER",
        },
        {
            "id": "synthetic-phone", "evidence_number": "SYN-JC-PHONE",
            "device_name": "SYNTHETIC HUAWEI SGU-AL10", "material_type": "phone",
            "material_type_status": "confirmed_by_user", "material_type_source": "user",
            "extractable": False,
        },
        {
            "id": "synthetic-iphone", "evidence_number": "SYN-JC-IPHONE",
            "device_name": "SYNTHETIC iPhone 14", "material_type": "phone",
            "material_type_status": "confirmed_by_user", "material_type_source": "user",
            "extractable": False,
        },
        {
            "evidence_number": "SYN-JC-NO-ID", "device_type": "SYNTHETIC vivo V2141A",
            "extractable": False, "imei1": "333333333333333",
            "serial_number": "SYNTHETIC-SERIAL-MUST-NOT-RENDER",
        },
    ]
    output = tmp_path / "material-contracts.docx"
    fill_template(report, str(_TEMPLATE), str(output))

    with zipfile.ZipFile(output) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")

    for expected in (
        "111111111111111",
        "TABLET-SERIAL-MUST-RENDER",
        "SYNTHETIC HUAWEI SGU-AL10手机一部（无法提取）",
        "SYNTHETIC iPhone 14手机一部（无法提取）",
        "SYNTHETIC vivo V2141A一部（无法提取）",
    ):
        assert expected in document_xml
    for forbidden in (
        "PHONE-SERIAL-MUST-NOT-RENDER",
        "222222222222222",
        "SYNTHETIC HUAWEI SGU-AL10一部",
        "SYNTHETIC iPhone 14一部",
        "333333333333333",
        "SYNTHETIC-SERIAL-MUST-NOT-RENDER",
    ):
        assert forbidden not in document_xml


def test_evidence_renderer_preserves_unconfirmed_template_name_priority(tmp_path):
    report = _report()
    report["introduction"]["evidence_list"] = [{
        "evidence_number": "SYN-JC-UNCONFIRMED",
        "device_type": "SYNTHETIC Android设备",
        "device_name": "SYNTHETIC HUAWEI",
        "model": "SYNTHETIC MODEL",
        "material_type": "unconfirmed",
        "material_type_status": "unconfirmed",
        "extractable": False,
    }]
    output = tmp_path / "unconfirmed-name-priority.docx"
    fill_template(report, str(_TEMPLATE), str(output))

    with zipfile.ZipFile(output) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
    assert "SYNTHETIC HUAWEI一部（无法提取）" in document_xml
    assert "SYNTHETIC Android设备一部" not in document_xml


def test_manifest_result_uses_every_part_filename_hash_size_and_disc(tmp_path):
    report = _report()
    report["introduction"]["evidence_list"] = [
        {"evidence_number": "JC-A", "device_type": "测试手机"},
        {"evidence_number": "JC-B", "device_type": "测试手机"},
        {"evidence_number": "JC-C", "device_type": "测试手机"},
    ]
    report["inspection"]["primary_software"] = {
        "name": "已确认取证软件",
        "version": "3.2",
        "confirmation_status": "confirmed_by_report",
    }
    report["inspection"]["software_tools"] = [
        {"name": "WinRAR压缩管理软件", "version": "6.24"},
        {"name": "Python hashlib", "version": "3.12"},
    ]
    output = tmp_path / "manifest-result.docx"
    fill_template(report, str(_TEMPLATE), str(output), [], build_archive_manifest())

    with zipfile.ZipFile(output) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")

    document = Document(output)
    title = next(p for p in document.paragraphs if p.text.strip() == "电子数据检查笔录")
    extract_heading = next(
        p for p in document.paragraphs if p.text.strip() == "电子数据提取固定清单"
    )
    assert title.alignment == 1
    assert title.runs and all(run.bold for run in title.runs if run.text)
    assert extract_heading.runs and all(run.bold for run in extract_heading.runs if run.text)

    assert "经对编号为JC-A、JC-B、JC-C号检材使用已确认取证软件（版本号为3.2）" in document_xml
    for index in range(1, 4):
        assert f"synthetic.part{index}.rar" in document_xml
        assert f"文件大小为“{index * 100}”字节" in document_xml
        assert f"GP20260706-{index:02d}" in document_xml
    assert "case.rar" not in document_xml
    assert "client-value.rar" not in document_xml


def test_fill_template_preserves_vml_and_renders_default_and_pagination(tmp_path):
    output = tmp_path / "filled.docx"
    fill_template(_report(), str(_TEMPLATE), str(output))

    with zipfile.ZipFile(output) as package:
        assert package.testzip() is None
        document_xml = package.read("word/document.xml").decode("utf-8")

    assert document_xml.count("<w:pict") >= 2
    assert document_xml.count("<v:textbox") == 2
    assert document_xml.count("<w:txbxContent") == 2
    assert "检验单位：测试鉴定中心" in document_xml
    assert _DEFAULT_SUMMARY in document_xml
    assert "<w:pageBreakBefore" not in document_xml
    assert document_xml.count('w:type="page"') == 3


def test_attachment_summary_uses_page_top_suppressed_spacing_and_conditional_pagination(tmp_path):
    report = _report()
    report["introduction"]["evidence_list"] = [
        {"evidence_number": "SYN-JC-001", "device_type": "测试手机"},
        {"evidence_number": "SYN-JC-002", "device_type": "测试平板"},
    ]
    report["inspection"]["primary_software"] = {
        "name": "SYNTHETIC 长名称取证软件",
        "version": "2026.08.12-TEST",
        "confirmation_status": "confirmed_by_user",
    }
    report["inspection"]["software_tools"] = [
        {"name": "WinRAR压缩管理软件", "version": "6.24"},
        {"name": "Python hashlib", "version": "3.12"},
    ]
    output = tmp_path / "long-result.docx"

    fill_template(report, str(_TEMPLATE), str(output), [], build_archive_manifest(2))

    with zipfile.ZipFile(output) as package:
        root = ET.fromstring(package.read("word/document.xml"))
    body = root.find("./{%s}body" % W_NS)
    children = list(body)
    summary_index = next(
        index for index, element in enumerate(children)
        if "1、电子数据提取固定清单" in "".join(element.itertext())
    )
    blank_count = 0
    cursor = summary_index - 1
    while cursor >= 0 and not "".join(children[cursor].itertext()).strip():
        blank_count += 1
        cursor -= 1
    previous_text = "".join(children[cursor].itertext()).strip()

    assert previous_text.startswith("经对编号为SYN-JC-001、SYN-JC-002号检材")
    assert blank_count == 0
    spacing = children[summary_index].find('./{%s}pPr/{%s}spacing' % (W_NS, W_NS))
    assert spacing is not None
    assert spacing.get('{%s}before' % W_NS) == '1560'
    assert len(children[summary_index].findall(
        './/{%s}br[@{%s}type="page"]' % (W_NS, W_NS)
    )) == 0

    summary_paragraphs = [
        element for element in children
        if any(marker in "".join(element.itertext()) for marker in (
            "1、电子数据提取固定清单",
            "2、检材图",
            "3、本鉴定中心刻制的",
        ))
    ]
    assert len(summary_paragraphs) == 3
    for paragraph in summary_paragraphs:
        assert paragraph.find('./{%s}pPr/{%s}keepLines' % (W_NS, W_NS)) is not None
    for paragraph in summary_paragraphs:
        assert paragraph.find('./{%s}pPr/{%s}keepNext' % (W_NS, W_NS)) is not None

    attachment1 = next(
        element for element in children
        if "".join(element.itertext()).strip() == "附件1："
    )
    attachment1_index = children.index(attachment1)
    summary_block = children[summary_index:attachment1_index]
    for paragraph in summary_block:
        assert paragraph.find('./{%s}pPr/{%s}keepLines' % (W_NS, W_NS)) is not None
    for paragraph in summary_block[:-1]:
        assert paragraph.find('./{%s}pPr/{%s}keepNext' % (W_NS, W_NS)) is not None
    assert summary_block[-1].find(
        './{%s}pPr/{%s}keepNext' % (W_NS, W_NS)
    ) is None
    assert len(attachment1.findall(
        './/{%s}br[@{%s}type="page"]' % (W_NS, W_NS)
    )) == 1


def test_attachment_summary_signature_and_date_layout_stay_unchanged(tmp_path):
    output = tmp_path / "summary-layout.docx"
    report = _report()
    report["inspection"]["primary_software"] = {
        "name": "SYNTHETIC 测试取证软件",
        "version": "1.0-TEST",
        "confirmation_status": "confirmed_by_report",
    }
    report["inspection"]["software_tools"] = [
        {"name": "WinRAR压缩管理软件", "version": "6.24"},
        {"name": "Python hashlib", "version": "3.12"},
    ]
    fill_template(report, str(_TEMPLATE), str(output), [], build_archive_manifest())

    def body_children(path):
        with zipfile.ZipFile(path) as package:
            root = ET.fromstring(package.read("word/document.xml"))
        return list(root.find("./{%s}body" % W_NS))

    template_children = body_children(_TEMPLATE)
    output_children = body_children(output)

    def layout_region(children):
        start = next(
            index for index, element in enumerate(children)
            if "1、电子数据提取固定清单" in "".join(element.itertext())
        )
        end = next(
            index for index, element in enumerate(children[start:], start)
            if "附件1：" in "".join(element.itertext())
        )
        return children[start:end]

    template_region = layout_region(template_children)
    output_region = layout_region(output_children)
    assert len(output_region) == len(template_region)

    def layout_property(element, local_name):
        value = element.find('./{%s}pPr/{%s}%s' % (W_NS, W_NS, local_name))
        return ET.tostring(value) if value is not None else None

    for index, (expected, actual) in enumerate(zip(template_region, output_region)):
        for local_name in ("ind", "jc", "tabs"):
            assert layout_property(actual, local_name) == layout_property(
                expected, local_name
            )
        expected_spacing = expected.find('./{%s}pPr/{%s}spacing' % (W_NS, W_NS))
        actual_spacing = actual.find('./{%s}pPr/{%s}spacing' % (W_NS, W_NS))
        assert expected_spacing is not None
        assert actual_spacing is not None
        expected_attributes = dict(expected_spacing.attrib)
        actual_attributes = dict(actual_spacing.attrib)
        if index == 0:
            assert actual_attributes.pop('{%s}before' % W_NS) == '1560'
        assert actual_attributes == expected_attributes
        assert [ET.tostring(node) for node in actual.findall('.//{%s}pict' % W_NS)] == [
            ET.tostring(node) for node in expected.findall('.//{%s}pict' % W_NS)
        ]


def test_generated_docx_removes_comments_and_nonessential_metadata(tmp_path):
    output = tmp_path / "sanitized.docx"
    fill_template(_report(), str(_TEMPLATE), str(output))

    with zipfile.ZipFile(output) as package:
        names = set(package.namelist())
        assert "word/comments.xml" not in names
        assert "word/commentsExtended.xml" not in names
        assert "word/people.xml" not in names
        assert "docProps/custom.xml" not in names
        xml_parts = {
            name: package.read(name).decode("utf-8")
            for name in names
            if name.endswith(".xml") or name.endswith(".rels")
        }
        document_xml = xml_parts["word/document.xml"]
        settings_xml = xml_parts["word/settings.xml"]
        core = ET.fromstring(xml_parts["docProps/core.xml"])
        content_types = xml_parts["[Content_Types].xml"]
        package_rels = xml_parts["_rels/.rels"]

    assert "commentRangeStart" not in document_xml
    assert "commentRangeEnd" not in document_xml
    assert "commentReference" not in document_xml
    assert "comments.xml" not in "".join(xml_parts.values())
    assert "commentsExtended.xml" not in "".join(xml_parts.values())
    assert "people.xml" not in "".join(xml_parts.values())
    assert "docVars" not in settings_xml
    assert "custom.xml" not in content_types
    assert "custom.xml" not in package_rels
    tracked_change_names = {"ins", "del", "moveFrom", "moveTo"}
    for xml in xml_parts.values():
        root = ET.fromstring(xml)
        assert not any(
            node.tag.rsplit("}", 1)[-1] in tracked_change_names
            for node in root.iter()
        )

    core_values = {
        node.tag.rsplit("}", 1)[-1]: node.text or ""
        for node in core
    }
    assert core_values["creator"] == "文枢"
    assert core_values["lastModifiedBy"] == "文枢"
    assert core_values["revision"] == "1"
    assert "lastPrinted" not in core_values
    assert "subject" not in core_values
    assert "keywords" not in core_values

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for name, xml in xml_parts.items():
        if not name.startswith("word/") or not name.endswith(".xml"):
            continue
        root = ET.fromstring(xml)
        for run in root.findall(".//{%s}r" % w_ns):
            color = run.find("./{%s}rPr/{%s}color" % (w_ns, w_ns))
            assert color is not None
            assert color.get("{%s}val" % w_ns) == "000000"


def _write_png(path: Path, width: int, height: int, color=(50, 120, 200)):
    """使用标准库生成可被 Word 读取的真实 PNG 样例。"""
    raw = b"".join(b"\x00" + bytes(color) * width for _ in range(height))

    def chunk(kind, data):
        return (struct.pack(">I", len(data)) + kind + data
                + struct.pack(">I", zlib.crc32(kind + data) & 0xFFFFFFFF))

    payload = b"\x89PNG\r\n\x1a\n"
    payload += chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
    payload += chunk(b"IDAT", zlib.compress(raw, 1))
    payload += chunk(b"IEND", b"")
    path.write_bytes(payload)


@pytest.mark.parametrize("sizes", [
    [],
    [(4000, 4000)],      # 1张超出页面可用区域的图
    [(1600, 900), (900, 1600), (1200, 1200), (2000, 1000)],  # 4张图片
])
def test_photo_regression_scenarios_keep_images_and_page_xml(tmp_path, sizes):
    photos = []
    for index, (width, height) in enumerate(sizes):
        path = tmp_path / f"photo-{index}.png"
        _write_png(path, width, height, color=(50 + index * 20, 120, 200))
        photos.append(str(path))

    output = tmp_path / "photos.docx"
    fill_template(_report(), str(_TEMPLATE), str(output), photos)

    with zipfile.ZipFile(output) as package:
        assert package.testzip() is None
        document_xml = package.read("word/document.xml").decode("utf-8")
        root = ET.fromstring(document_xml)

    assert document_xml.count("<w:drawing") == len(photos)
    doc_pr_ids = [doc_pr.get("id") for doc_pr in root.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr")]
    assert len(doc_pr_ids) == len(set(doc_pr_ids))
    extents = [
        (int(extent.get("cx")), int(extent.get("cy")))
        for extent in root.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent")
    ]
    assert extents == [
        (
            calculate_fixed_geometry(width, height).render_width_emu,
            calculate_fixed_geometry(width, height).render_height_emu,
        )
        for width, height in sizes
    ]
    assert document_xml.count('w:type="page"') == 3
    assert 'w:type="oddPage"' not in document_xml
    assert 'w:type="evenPage"' not in document_xml
    assert "w:pageBreakBefore" not in document_xml
    assert "<w:keepNext" in document_xml
    assert "<w:keepLines" in document_xml


def test_report_only_export_keeps_three_material_photo_groups(tmp_path):
    report = _report()
    evidence_list = [
        {"id": f"material-{index}", "evidence_number": f"JC-{letter}", "device_type": "合成设备"}
        for index, letter in enumerate(("A", "B", "C"), 1)
    ]
    photo_ids = [f"photo-{index}" for index in range(1, 7)]
    report["introduction"]["evidence_list"] = evidence_list
    report["attachments"]["photo_ids"] = photo_ids
    report["attachments"]["photo_groups"] = [
        {
            "material_id": item["id"],
            "material_number": item["evidence_number"],
            "display_text": f"检材{item['evidence_number']}照片",
            "ordered_image_ids": photo_ids[index * 2:index * 2 + 2],
            "source_order": index + 1,
        }
        for index, item in enumerate(evidence_list)
    ]
    photo_paths = []
    sizes = [
        (1600, 400), (400, 1600), (1000, 1000),
        (1200, 600), (600, 1200), (800, 800),
    ]
    for index, (width, height) in enumerate(sizes, 1):
        path = tmp_path / f"SYNTHETIC-photo-{index}.png"
        _write_png(path, width, height)
        photo_paths.append(str(path))

    output = tmp_path / "report-only-three-materials.docx"
    fill_template(report, str(_TEMPLATE), str(output), photo_paths)

    root = ET.fromstring(zipfile.ZipFile(output).read("word/document.xml"))
    body = root.find("./{%s}body" % W_NS)
    tables = [
        table for table in body.findall("./{%s}tbl" % W_NS)
        if "检材JC-" in "".join(table.itertext())
    ]
    assert len(tables) == 2
    assert [
        [len(row.findall("./{%s}tc" % W_NS)) for row in table.findall("./{%s}tr" % W_NS)]
        for table in tables
    ] == [[2, 1, 1, 2, 1], [2, 1]]
    page_breaks = [list(body)[list(body).index(table) - 1] for table in tables]
    profile = current_template_profile()
    assert ["".join(page.itertext()).strip() for page in page_breaks] == ["附件2：", ""]
    assert [
        page.find("./{%s}pPr/{%s}spacing" % (W_NS, W_NS)).get("{%s}after" % W_NS)
        for page in page_breaks
    ] == ["0", str(profile.attachment2_page_break_after_twips)]
    extents = [
        (int(extent.get("cx")), int(extent.get("cy")))
        for extent in root.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent"
        )
    ]
    assert extents == [
        (
            calculate_fixed_geometry(
                width, height,
                slot_height_emu=(
                    ATTACHMENT2_DUAL_GROUP_SLOT_HEIGHT_EMU
                    if index < 4 else ATTACHMENT2_SLOT_HEIGHT_EMU
                ),
            ).render_width_emu,
            calculate_fixed_geometry(
                width, height,
                slot_height_emu=(
                    ATTACHMENT2_DUAL_GROUP_SLOT_HEIGHT_EMU
                    if index < 4 else ATTACHMENT2_SLOT_HEIGHT_EMU
                ),
            ).render_height_emu,
        )
        for index, (width, height) in enumerate(sizes)
    ]
    text = "".join(root.itertext())
    assert [text.count(f"检材JC-{letter}照片") for letter in ("A", "B", "C")] == [1, 1, 1]
