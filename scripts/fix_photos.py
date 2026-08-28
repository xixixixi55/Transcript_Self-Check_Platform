#!/usr/bin/env python3
"""修复 _handle_photos：标题位于图片下方，并针对 1 张与 2 张以上照片自适应布局。"""
import re, os

path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                    'packages', 'backend', 'app', 'services', 'template_filler_service.py')
with open(path, 'r', encoding='utf-8') as f:
    content = f.read()

old_func_start = content.find('def _handle_photos(doc: Document, photo_paths: list[str], report: dict):')
old_func_end = content.find('\n\n# ═', old_func_start)
assert old_func_start > 0 and old_func_end > 0, f"Function not found: {old_func_start}, {old_func_end}"

new_func = r'''def _handle_photos(doc: Document, photo_paths: list[str], report: dict):
    """处理附件2 检材照片 — 标题在图片下方，自适应布局"""
    from lxml import etree
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    v_imagedata = '{urn:schemas-microsoft-com:vml}imagedata'
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body

    # 1. 删除模板图片段落
    to_remove = []
    for pe in body.findall('.//{' + w_ns + '}p'):
        if pe.findall('.//' + v_imagedata):
            to_remove.append(pe)
    for pe in to_remove:
        body.remove(pe)

    # 2. 无照片：清空标题
    if not photo_paths:
        for para in doc.paragraphs:
            if "检材" in para.text and "照片" in para.text:
                for run in para.runs:
                    run.text = ""
        return

    # 3. 找到照片标题段落（图片将插在标题之前，使标题在图片下方）
    caption_idx = None
    for pi, para in enumerate(doc.paragraphs):
        if "检材" in para.text and "照片" in para.text:
            caption_idx = pi
            break
    if caption_idx is None:
        return

    caption_elem = doc.paragraphs[caption_idx]._element

    # 4. 根据照片数量选择布局
    if len(photo_paths) == 1:
        img_el = _make_single_image(doc, photo_paths[0], w_ns)
        if img_el is not None:
            caption_elem.addprevious(img_el)
    else:
        tbl = _make_image_grid(doc, photo_paths, w_ns)
        if tbl is not None:
            caption_elem.addprevious(tbl)

    # 5. 更新标题文本中的检材编号
    intro = report.get("introduction", {})
    ev_list = intro.get("evidence_list", [])
    evidence_num = ev_list[0].get("evidence_number", "") if ev_list else ""
    for para in doc.paragraphs:
        if "检材" in para.text and "照片" in para.text:
            for run in para.runs:
                if "{{first_evidence_number}}" in run.text:
                    run.text = run.text.replace("{{first_evidence_number}}", evidence_num or "")
                elif "first_evidence_number" in run.text:
                    run.text = run.text.replace("first_evidence_number", evidence_num or "")
            break


def _make_single_image(doc, photo_path, w_ns):
    """创建居中单张图片段落"""
    from lxml import etree
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    if not os.path.isfile(photo_path):
        return None
    try:
        tmp_doc = Document()
        tmp_para = tmp_doc.add_paragraph()
        tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tmp_run = tmp_para.add_run()
        tmp_run.add_picture(photo_path, width=Inches(3.5), height=Inches(4.67))
        tmp_drawing = tmp_para._element.find('.//{' + w_ns + '}drawing')
        if tmp_drawing is None:
            return None
        img_part = doc.part.package.get_or_add_image_part(photo_path)
        rid = doc.part.relate_to(img_part, RT.IMAGE)
        for blip in tmp_drawing.findall(
                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            blip.set(
                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed',
                rid)
        new_p = etree.Element('{' + w_ns + '}p')
        new_pPr = etree.SubElement(new_p, '{' + w_ns + '}pPr')
        etree.SubElement(new_pPr, '{' + w_ns + '}jc').set('{' + w_ns + '}val', 'center')
        new_r = etree.SubElement(new_p, '{' + w_ns + '}r')
        new_r.append(tmp_drawing)
        return new_p
    except Exception:
        return None


def _make_image_grid(doc, photo_paths, w_ns):
    """创建 2 列无边框图片表格（2+ 张照片）"""
    from lxml import etree
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    cols = 2
    rows = (len(photo_paths) + cols - 1) // cols
    tbl = etree.Element('{' + w_ns + '}tbl')
    tblPr = etree.SubElement(tbl, '{' + w_ns + '}tblPr')
    tblW = etree.SubElement(tblPr, '{' + w_ns + '}tblW')
    tblW.set('{' + w_ns + '}w', '5000')
    tblW.set('{' + w_ns + '}type', 'pct')
    tblBorders = etree.SubElement(tblPr, '{' + w_ns + '}tblBorders')
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(tblBorders, '{' + w_ns + '}' + bn)
        b.set('{' + w_ns + '}val', 'none')
        b.set('{' + w_ns + '}sz', '0')
        b.set('{' + w_ns + '}space', '0')
        b.set('{' + w_ns + '}color', 'auto')
    tblGrid = etree.SubElement(tbl, '{' + w_ns + '}tblGrid')
    for c in range(cols):
        gc = etree.SubElement(tblGrid, '{' + w_ns + '}gridCol')
        gc.set('{' + w_ns + '}w', '4500')
    for ri in range(rows):
        tr = etree.SubElement(tbl, '{' + w_ns + '}tr')
        for ci in range(cols):
            pi = ri * cols + ci
            tc = etree.SubElement(tr, '{' + w_ns + '}tc')
            tcPr = etree.SubElement(tc, '{' + w_ns + '}tcPr')
            tcW = etree.SubElement(tcPr, '{' + w_ns + '}tcW')
            tcW.set('{' + w_ns + '}w', '4500')
            tcW.set('{' + w_ns + '}type', 'dxa')
            if pi < len(photo_paths) and os.path.isfile(photo_paths[pi]):
                try:
                    tmp_doc = Document()
                    tmp_para = tmp_doc.add_paragraph()
                    tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    tmp_run = tmp_para.add_run()
                    tmp_run.add_picture(photo_paths[pi], width=Inches(2.5), height=Inches(3.33))
                    tmp_drawing = tmp_para._element.find('.//{' + w_ns + '}drawing')
                    if tmp_drawing is not None:
                        img_part = doc.part.package.get_or_add_image_part(photo_paths[pi])
                        rid = doc.part.relate_to(img_part, RT.IMAGE)
                        for blip in tmp_drawing.findall(
                                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                            blip.set(
                                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed',
                                rid)
                        p = etree.SubElement(tc, '{' + w_ns + '}p')
                        pPr = etree.SubElement(p, '{' + w_ns + '}pPr')
                        etree.SubElement(pPr, '{' + w_ns + '}jc').set(
                            '{' + w_ns + '}val', 'center')
                        r = etree.SubElement(p, '{' + w_ns + '}r')
                        r.append(tmp_drawing)
                except Exception:
                    etree.SubElement(tc, '{' + w_ns + '}p')
            else:
                etree.SubElement(tc, '{' + w_ns + '}p')
    return tbl
'''

content = content[:old_func_start] + new_func + content[old_func_end:]
with open(path, 'w', encoding='utf-8') as f:
    f.write(content)
print('_handle_photos replaced successfully')
